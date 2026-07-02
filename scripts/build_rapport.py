# -*- coding: utf-8 -*-
"""
Final Year Project Report Generator - JobzFactory
Format: DOCX (Word) - compliant with the provided layout rules.
Structure mirrors the example Rapport_Hind_ABDELLAOUI_V2.pdf (Agile / SCRUM).

Expanded edition: aims at ~100 pages with richer technical content,
an added "Global Design and Modeling" chapter, fuller sprints and
extra appendices. Keeps the original styling exactly.
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

try:
    from PIL import Image as _PILImage
except Exception:
    _PILImage = None

OUT = os.path.join(os.path.dirname(__file__), "..", "Report", "JobzFactory_Rapport.docx")
OUT = os.path.abspath(OUT)

# Folder where the generated diagram / screenshot PNGs live.
ASSETS_DIR = os.path.join(os.path.dirname(__file__), "report_assets")


def _build_figure_map():
    """Map figure number -> asset filename.
    Figure counter order in the document:
      1  global use case            -> fig_01
      2  architecture               -> fig_02
      3  deployment                 -> fig_03
      4..8  tool logos (no image -> placeholder)
      9  global class diagram       -> fig_08
      10 global use case (design)   -> fig_09
      11 package diagram            -> fig_10
    Per sprint (6 figures: uc, class, seq1, seq2, iface1, iface2), sprint k=1..7:
      figure numbers start at 12 + (k-1)*6
      diagram files (uc,class,seq1,seq2) = fig_(11 + (k-1)*4) .. +3
      interface screenshots mapped explicitly below.
    Appendix screenshots: 54..59.
    """
    m = {1: "fig_01.png", 2: "fig_02.png", 3: "fig_03.png",
         9: "fig_08.png", 10: "fig_09.png", 11: "fig_10.png"}
    # sprint diagrams: 4 per sprint, file numbering sequential from fig_11
    for k in range(7):          # 0..6 -> sprint 1..7
        fig_base = 12 + k * 6   # first figure number of the sprint
        file_base = 11 + k * 4  # first diagram file of the sprint
        for j in range(4):      # uc, class, seq1, seq2
            m[fig_base + j] = "fig_%02d.png" % (file_base + j)
    # interface screenshots (iface1, iface2 = sprint base + 4, + 5)
    iface = {
        16: "fig_39.png", 17: "fig_40.png",   # sprint 1: register, profile
        22: "fig_41.png", 23: "fig_42.png",   # sprint 2: home, job detail
        28: "fig_43.png", 29: "fig_44.png",   # sprint 3: recruiter offers, board
        34: "fig_45.png", 35: "fig_46.png",   # sprint 4: admin offers, dashboard
        40: "fig_47.png", 41: "fig_48.png",   # sprint 5: apply, candidatures
        46: "fig_49.png", 47: "fig_50.png",   # sprint 6: search, pagination
        52: "fig_51.png", 53: "fig_46.png",   # sprint 7: recruiter dashboard, admin dashboard
        54: "fig_41.png", 55: "fig_42.png",   # appendix: home, job detail
        56: "fig_51.png", 57: "fig_45.png",   # appendix: recruiter dashboard, admin offers
        58: "fig_47.png", 59: "fig_49.png",   # appendix: apply, search
    }
    m.update(iface)
    return m


FIG_IMG_MAP = _build_figure_map()


def _figure_image_path(fig_number):
    fname = FIG_IMG_MAP.get(fig_number)
    if not fname:
        return None
    p = os.path.join(ASSETS_DIR, fname)
    return p if os.path.exists(p) else None

NAVY = RGBColor(0x1A, 0x44, 0x80)
BROWN = RGBColor(0xA0, 0x4A, 0x1E)
CAPBLUE = RGBColor(0x2E, 0x74, 0xB5)
GREY = RGBColor(0x59, 0x59, 0x59)

# ----------------------------- counters -----------------------------
FIG = 0          # figure counter
TAB = 0          # table counter
FIG_LIST = []    # collected "Figure N: title"
TAB_LIST = []    # collected "Table N: title"

# ----------------------------- helpers -----------------------------

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            data = kwargs[edge]
            el = tcBorders.find(qn('w:' + edge))
            if el is None:
                el = OxmlElement('w:' + edge)
                tcBorders.append(el)
            el.set(qn('w:val'), data.get('val', 'single'))
            el.set(qn('w:sz'), str(data.get('sz', 4)))
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), data.get('color', 'BFBFBF'))


def shade_cell(cell, fill="F2F2F2"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def field(run, instr):
    """Insert a Word field (TOC, PAGE, etc.) into a run."""
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin')
    i = OxmlElement('w:instrText'); i.set(qn('xml:space'), 'preserve'); i.text = instr
    s = OxmlElement('w:fldChar'); s.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = ""
    e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end')
    run._r.append(b); run._r.append(i); run._r.append(s); run._r.append(t); run._r.append(e)


def style_normal(doc):
    st = doc.styles['Normal']
    st.font.name = 'Times New Roman'
    st.font.size = Pt(12)
    pf = st.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_after = Pt(6)


def setup_headings(doc):
    for name, size in (('Heading 1', 16), ('Heading 2', 14), ('Heading 3', 12)):
        st = doc.styles[name]
        st.font.name = 'Times New Roman'
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = NAVY
        st.paragraph_format.space_before = Pt(12)
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.keep_with_next = True


def add_para(doc, text, style=None, align=None, bold=False, italic=False, color=None, size=None, space_after=6):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if color is not None:
        r.font.color.rgb = color
    if size is not None:
        r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_body(doc, text):
    return add_para(doc, text)


def h1(doc, text):
    return add_para(doc, text, style='Heading 1')


def h2(doc, text):
    return add_para(doc, text, style='Heading 2')


def h3(doc, text):
    return add_para(doc, text, style='Heading 3')


def h3brown(doc, text):
    p = add_para(doc, text, style='Heading 3')
    for r in p.runs:
        r.font.color.rgb = BROWN
    return p


def _insert_figure_image(doc, image_path):
    """Insert a centered PNG, width ~15cm but capped at ~17cm height so tall
    diagrams still fit on a page. Returns the picture paragraph."""
    target_w = Cm(15)
    max_h = Cm(17)
    width = target_w
    height = None
    if _PILImage is not None:
        try:
            with _PILImage.open(image_path) as im:
                w, h = im.size
            if h > 0 and w > 0:
                aspect = h / w
                # height in cm at target width
                height_cm = 15.0 * aspect
                if height_cm > 17.0:
                    # cap by height, let width shrink
                    height = max_h
                    width = None
        except Exception:
            pass
    if width is not None:
        doc.add_picture(image_path, width=width)
    else:
        doc.add_picture(image_path, height=height)
    pic_para = doc.paragraphs[-1]
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # tighten spacing around the picture
    pic_para.paragraph_format.space_before = Pt(2)
    pic_para.paragraph_format.space_after = Pt(2)
    return pic_para


def add_figure(doc, title, description, box_h=Cm(7)):
    """Figure with caption. If a generated PNG exists for this figure number
    (scripts/report_assets/fig_NN.png) it is inserted as a real image;
    otherwise a bordered placeholder box is drawn. The blue italic caption and
    the auto-incrementing counter behave identically in both cases."""
    global FIG
    FIG += 1
    FIG_LIST.append("Figure %d: %s" % (FIG, title))
    img = _figure_image_path(FIG)
    if img:
        _insert_figure_image(doc, img)
    else:
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        cell = tbl.cell(0, 0)
        cell.width = Cm(15)
        # force the row height
        tr = tbl.rows[0]._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), str(int(box_h.emu / 635)))  # EMU -> twips
        trHeight.set(qn('w:hRule'), 'atLeast')
        trPr.append(trHeight)
        set_cell_border(cell, top={'sz': 8, 'color': 'BFBFBF'}, left={'sz': 8, 'color': 'BFBFBF'},
                        bottom={'sz': 8, 'color': 'BFBFBF'}, right={'sz': 8, 'color': 'BFBFBF'})
        shade_cell(cell, "F7F9FC")
        cp = cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cp.add_run(description)
        run.italic = True
        run.font.color.rgb = GREY
        run.font.size = Pt(10)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run("Figure %d: %s" % (FIG, title))
    cr.italic = True
    cr.font.color.rgb = CAPBLUE
    cr.font.size = Pt(11)
    cap.paragraph_format.space_after = Pt(10)
    return None


def add_table_caption(doc, title):
    """Auto-incrementing table caption."""
    global TAB
    TAB += 1
    TAB_LIST.append("Table %d: %s" % (TAB, title))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run("Table %d: %s" % (TAB, title))
    cr.italic = True
    cr.font.color.rgb = CAPBLUE
    cr.font.size = Pt(11)
    cap.paragraph_format.space_after = Pt(4)
    cap.paragraph_format.space_before = Pt(6)


def make_table(doc, headers, rows, widths=None, font_size=10):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = hdr[i].paragraphs[0].add_run(htext)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(hdr[i], "1A4480")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].paragraphs[0].add_run(str(val)).font.size = Pt(font_size)
            cells[i].paragraphs[0].paragraph_format.space_after = Pt(2)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = w
    return t


def add_page_number_footer(section):
    section.different_first_page_header_footer = True
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Page | ")
    r.font.size = Pt(10)
    r.font.color.rgb = GREY
    field(r, "PAGE")


def set_pgnum_start(section, start):
    sectPr = section._sectPr
    pn = sectPr.find(qn('w:pgNumType'))
    if pn is None:
        pn = OxmlElement('w:pgNumType')
        sectPr.append(pn)
    pn.set(qn('w:start'), str(start))


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin')
    i = OxmlElement('w:instrText'); i.set(qn('xml:space'), 'preserve')
    i.text = 'TOC \\o "1-3" \\h \\z \\u'
    s = OxmlElement('w:fldChar'); s.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = "Table of contents: in Word, right-click here > Update fields (F9)."
    e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end')
    run._r.append(b); run._r.append(i); run._r.append(s); run._r.append(t); run._r.append(e)


def hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), 'BFBFBF')
    pbdr.append(bottom); pPr.append(pbdr)
    p.paragraph_format.space_after = Pt(8)


def chapter_title(doc, text):
    h1(doc, text)
    hr(doc)


def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text)
    p.paragraph_format.space_after = Pt(3)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.add_run(text)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_code(doc, code, caption=None):
    """Monospaced code block in a single bordered cell (Consolas 9pt)."""
    if caption:
        cp = doc.add_paragraph()
        cr = cp.add_run(caption)
        cr.italic = True
        cr.font.color.rgb = CAPBLUE
        cr.font.size = Pt(11)
        cp.paragraph_format.space_after = Pt(2)
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    cell.width = Cm(16)
    set_cell_border(cell, top={'sz': 6, 'color': 'BFBFBF'}, left={'sz': 6, 'color': 'BFBFBF'},
                    bottom={'sz': 6, 'color': 'BFBFBF'}, right={'sz': 6, 'color': 'BFBFBF'})
    shade_cell(cell, "F4F4F4")
    first = True
    for line in code.split("\n"):
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(line if line else " ")
        r.font.name = 'Consolas'
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return tbl


def fill_list_before(marker_para, items):
    """Insert plain caption paragraphs before the marker paragraph element."""
    for text in items:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(3)
        marker_para._p.addprevious(p._p)


# ----------------------------- document -----------------------------

doc = Document()
style_normal(doc)
setup_headings(doc)

sec = doc.sections[0]
sec.page_height = Cm(29.7)
sec.page_width = Cm(21.0)
sec.top_margin = Cm(2.5)
sec.bottom_margin = Cm(2.5)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)
add_page_number_footer(sec)
set_pgnum_start(sec, 0)  # cover = page 0 (not displayed), Acknowledgements = Page | 1

# ===================== COVER PAGE =====================
def cover_page(doc):
    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, label in enumerate(["[Institution Logo]", "[Program Logo]", "[Company Logo]"]):
        c = tbl.cell(0, i)
        c.width = Cm(5)
        set_cell_border(c, top={'sz': 6, 'color': 'BFBFBF'}, left={'sz': 6, 'color': 'BFBFBF'},
                        bottom={'sz': 6, 'color': 'BFBFBF'}, right={'sz': 6, 'color': 'BFBFBF'})
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label); r.font.size = Pt(9); r.font.color.rgb = GREY; r.italic = True
    for _ in range(3):
        add_para(doc, "", space_after=4)

    add_para(doc, "Final Year Project Thesis", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14)
    add_para(doc, "Presented by", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    add_para(doc, "Oussama FISLY", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=18, color=NAVY, space_after=10)
    add_para(doc, "In partial fulfillment of the requirements for the degree of:", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    add_para(doc, "Professional Master's Degree", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13)
    add_para(doc, "Software Engineering (GL)", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13, space_after=18)

    add_para(doc, "Theme / Subject", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=12, color=GREY)
    hr(doc)
    add_para(doc,
             "Study and Development of a Web-Based Recruitment and Application Management Platform - JobzFactory",
             align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16, color=NAVY, space_after=6)
    hr(doc)
    for _ in range(6):
        add_para(doc, "", space_after=4)
    add_para(doc, "Supervised by:", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    add_para(doc, "Mr. Mohamed-Amine HAMMA", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13, color=NAVY, space_after=18)
    for _ in range(4):
        add_para(doc, "", space_after=4)
    add_para(doc, "Academic Year: 2025/2026", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12)

    doc.add_page_break()


cover_page(doc)

# ===================== ACKNOWLEDGEMENTS =====================
chapter_title(doc, "Acknowledgements")
add_body(doc,
    "First and foremost, I thank God Almighty for granting me the health, patience and willpower "
    "necessary to carry out this final year project under the best possible conditions.")
add_body(doc,
    "I would like to express my deepest gratitude to my supervisor, Mr. Mohamed-Amine HAMMA, "
    "for his availability, his insightful advice, his scientific rigor and his constant follow-up throughout this work. "
    "His guidance allowed me to make methodical progress and to achieve the set objectives.")
add_body(doc,
    "My thanks also go to the entire faculty of the Professional Master's degree in Software Engineering "
    "for the quality of the teaching I received, which provided me with the theoretical and practical foundations "
    "essential to the completion of this project.")
add_body(doc,
    "I also thank the members of the jury for the honor they do me in accepting to evaluate this work and for the time "
    "they will devote to its reading and evaluation.")
add_body(doc,
    "Finally, I dedicate this work to my parents and my family for their undeniable moral support, as well as to my friends "
    "and classmates for the moments shared throughout this academic journey.")
doc.add_page_break()

# ===================== ABSTRACT =====================
chapter_title(doc, "Abstract")
add_body(doc,
    "This report describes the work carried out as part of our final year project in the Professional Master's degree in "
    "Software Engineering. The project consists in designing and developing JobzFactory, a web-based recruitment platform "
    "aimed at connecting job seekers and recruiters. The platform provides candidates with a profile management space, "
    "an online job offer browsing and application system with CV upload, while offering recruiters a dedicated space for "
    "managing job advertisements and tracking applications. A central administration module handles the moderation of "
    "offers and recruiters.")
add_body(doc,
    "The project was developed following an agile SCRUM methodology, split into two releases grouping seven sprints. Each "
    "sprint followed a complete cycle of analysis, design, implementation and validation. The technical architecture relies "
    "on the ASP.NET MVC framework, the C# language, a SQL Server database for data storage and IIS for hosting. The user "
    "interface was built with HTML5, CSS3, JavaScript and Bootstrap to ensure a modern, responsive rendering.")
add_body(doc,
    "Keywords: Recruitment platform, ASP.NET MVC, C#, SQL Server, SCRUM, Agile, Job board, Application management, "
    "CV, Recruiter, Administration.")
doc.add_page_break()

# ===================== LIST OF ABBREVIATIONS =====================
chapter_title(doc, "List of Abbreviations")
abbr = [
    ("API", "Application Programming Interface"),
    ("ADO.NET", "ActiveX Data Objects for .NET"),
    ("ASP.NET", "Active Server Pages .NET"),
    ("BLL", "Business Logic Layer"),
    ("CLR", "Common Language Runtime"),
    ("CSS", "Cascading Style Sheets"),
    ("CV", "Curriculum Vitae"),
    ("DAL", "Data Access Layer"),
    ("FCL", "Framework Class Library"),
    ("HTML", "HyperText Markup Language"),
    ("HTTP", "HyperText Transfer Protocol"),
    ("IIS", "Internet Information Services"),
    ("IL", "Intermediate Language"),
    ("JS", "JavaScript"),
    ("JIT", "Just-In-Time compiler"),
    ("MVC", "Model-View-Controller"),
    ("ORM", "Object-Relational Mapping"),
    ("PDF", "Portable Document Format"),
    ("PFE", "Final Year Project"),
    ("DBMS", "Database Management System"),
    ("SQL", "Structured Query Language"),
    ("SSMS", "SQL Server Management Studio"),
    ("T-SQL", "Transact-SQL"),
    ("UML", "Unified Modeling Language"),
    ("URL", "Uniform Resource Locator"),
    ("UI", "User Interface"),
    ("UX", "User Experience"),
    ("XML", "eXtensible Markup Language"),
]
make_table(doc, ["Abbreviation", "Meaning"], abbr, widths=[Cm(4), Cm(11)])
doc.add_page_break()

# ===================== LIST OF FIGURES =====================
chapter_title(doc, "List of Figures")
lof_marker = doc.add_paragraph()  # filled at the end
doc.add_page_break()

# ===================== LIST OF TABLES =====================
chapter_title(doc, "List of Tables")
lot_marker = doc.add_paragraph()  # filled at the end
doc.add_page_break()

# ===================== TABLE OF CONTENTS =====================
chapter_title(doc, "Table of Contents")
add_toc(doc)
doc.add_page_break()

# ===================== GENERAL INTRODUCTION =====================
chapter_title(doc, "General Introduction")
add_body(doc,
    "The rise of information and communication technologies has profoundly transformed the world of work and, in "
    "particular, the recruitment process. Today, almost all companies use digital platforms to publish their job offers "
    "and collect applications. This mode of recruitment, while faster and broader in reach, nevertheless presents "
    "challenges: a large volume of offers to moderate, a diversity of profiles, the need for rigorous application tracking "
    "and a smooth user experience for both candidates and recruiters.")
add_body(doc,
    "JobzFactory is a web-based recruitment platform whose objective is to provide a centralized, simple and efficient "
    "solution to connect job seekers and recruiters. On the candidate side, the platform allows the creation of a profile, "
    "the browsing of job offers, advanced search and online application with a PDF CV upload. On the recruiter side, it "
    "offers a dedicated space for publishing advertisements, tracking applications and managing the history of offers. "
    "Finally, an administration module makes it possible to supervise all content, manage recruiters and moderate offers.")
add_body(doc,
    "This report describes all the work carried out as part of this final year project. The first chapter presents the "
    "general context of the project, its problem statement, its objectives, the analysis of existing solutions and the "
    "adopted development methodology. The second chapter is devoted to the analysis of functional and non-functional "
    "requirements, the identification of actors and use cases, and the description of business processes. The third "
    "chapter details the technical study, the architecture choices and the tools used. The fourth chapter presents the "
    "global design and modeling of the platform. Chapters five and six present the implementation of the project following "
    "an agile SCRUM approach, broken down into seven sprints grouped into two releases. Each sprint is documented "
    "according to the same pattern: description, analysis and design, implementation and conclusion. The report ends with "
    "a general conclusion, a glossary, a bibliography, a webography and appendices.")
doc.add_page_break()

# ===================== CHAPTER 1 =====================
chapter_title(doc, "CHAPTER 1: General Project Context")
h2(doc, "1. Project context")
h3(doc, "1.1. General problem statement")
add_body(doc,
    "Traditional recruitment, based on publishing press advertisements and receiving paper applications, proves to be slow, "
    "costly and limited in reach. With dematerialization, companies are turning to online platforms that make it possible "
    "to reach a wider audience and to centralize applications. However, many existing platforms suffer from a fragmented "
    "user experience: poorly ergonomic interfaces, lack of a dedicated space for recruiters, difficult application tracking, "
    "and a lack of moderation and administration tools.")
add_body(doc,
    "The problem statement of our project can be formulated as follows: how to design and develop a unified web-based "
    "recruitment platform that offers a smooth experience to candidates and recruiters, ensures rigorous application "
    "tracking and provides a back-office of administration allowing the moderation and the steering of all content? "
    "This problematic is declined into three sub-questions relative to the candidate experience, the recruiter experience "
    "and the administration and moderation of the published content.")

h3(doc, "1.2. Context & relevance")
add_body(doc,
    "The project is set in a context where the online recruitment market is growing steadily. Job board type platforms "
    "have become an essential channel for companies and candidates. The relevance of the project is twofold. On the one "
    "hand, it responds to a concrete market need by proposing a complete and coherent solution. On the other hand, "
    "academically, it constitutes a rich application ground allowing the implementation of the skills acquired during the "
    "Software Engineering Master's: software design, web programming, database design, agile project management and "
    "software quality.")

h3(doc, "1.3. Project goal")
add_body(doc,
    "The goal of the project is to design, develop and deliver an operational web-based recruitment platform, named "
    "JobzFactory, allowing: (1) candidates to create a profile, browse and search for offers, and apply online; (2) "
    "recruiters to publish advertisements, manage their offers and consult the applications received; (3) administrators "
    "to moderate offers, validate recruiters and steer the platform through a dashboard.")

h3(doc, "1.4. Project missions")
for m in [
    "Identify and analyze the functional and non-functional requirements of the three categories of users.",
    "Design the application architecture and the data architecture of the platform.",
    "Develop the various modules following an agile SCRUM approach.",
    "Build responsive and ergonomic user interfaces.",
    "Set up an administration and statistics module.",
    "Test the features and validate compliance with the expressed requirements.",
    "Produce the UML modeling and the technical and user documentation.",
]:
    bullet(doc, m)

h3(doc, "1.5. Project deliverables")
for l in [
    "This final year project report.",
    "The application source code (multi-project Visual Studio solution).",
    "The database creation and initialization script.",
    "The user manual and the technical documentation.",
    "The UML diagrams (use case, class, sequence, deployment).",
    "The deployment and administration guide.",
]:
    bullet(doc, l)

h3(doc, "1.6. Benchmark of existing solutions")
add_body(doc,
    "Before designing JobzFactory, we carried out a benchmark of the main existing online recruitment solutions in order "
    "to identify their strengths and weaknesses and to position our platform relative to them. The compared solutions "
    "include generalist international job boards as well as regional platforms. The comparison criteria retained are the "
    "candidate experience, the recruiter space, the moderation, the search capabilities and the cost.")
add_table_caption(doc, "Benchmark of existing recruitment solutions")
make_table(doc, ["Solution", "Candidate experience", "Recruiter space", "Moderation", "Search", "Cost"],
    [
        ["LinkedIn Jobs", "Strong profile, network", "Advanced, paid", "Self-published", "Keyword + filters", "High"],
        ["Indeed", "Easy CV upload", "Sponsored offers", "Light", "Keyword + location", "Medium"],
        ["Glassdoor", "Reviews + jobs", "Employer branding", "Light", "Keyword + filters", "Medium"],
        ["Rekrute", "Regional reach", "Dedicated space", "Manual", "Keyword + sector", "Medium"],
        ["Anapec Jobs", "Public service", "Basic", "Manual", "Basic filters", "Free"],
        ["JobzFactory (target)", "Profile + CV + apply", "Full CRUD + tracking", "Admin back-office", "Advanced + pagination", "Free (academic)"],
    ], widths=[Cm(3.2), Cm(3.0), Cm(3.0), Cm(2.2), Cm(2.6), Cm(2.0)])

h3(doc, "1.7. SWOT analysis")
add_body(doc,
    "A SWOT analysis was carried out to position the project with respect to its environment. It highlights the internal "
    "strengths and weaknesses of the platform as well as the external opportunities and threats that may influence its "
    "evolution.")
add_table_caption(doc, "SWOT analysis of the JobzFactory project")
make_table(doc, ["Strengths", "Weaknesses", "Opportunities", "Threats"],
    [
        ["Unified 3-in-1 platform (candidate / recruiter / admin)",
         "Academic scope, limited volume of test data",
         "Growing digitalization of recruitment in the region",
         "Strong competition from established international job boards"],
        ["Modern stack (ASP.NET MVC, SQL Server, Bootstrap)",
         "No matching engine yet",
         "Possible integration with national employment programs",
         "Changes in legal framework on personal data"],
        ["Agile delivery, evolvable architecture",
         "Single developer, limited testing resources",
         "Extension to a mobile application",
         "User retention facing larger networks"],
        ["Moderation back-office ensuring content quality",
         "Manual moderation, not yet automated",
         "B2B partnerships with local companies",
         "Cybersecurity risks (CV data, credentials)"],
    ], widths=[Cm(4), Cm(4), Cm(4), Cm(4)])

h3(doc, "1.8. Project scope")
add_body(doc,
    "Defining the scope of the project makes it possible to clearly distinguish what is part of the deliverable from what "
    "is explicitly excluded, in order to avoid scope drift during the sprints.")
add_table_caption(doc, "Project scope (in / out of scope)")
make_table(doc, ["In scope", "Out of scope (future work)"],
    [
        ["Candidate registration, authentication, profile and CV management",
         "Automated candidate-to-offer matching engine"],
        ["Publication, editing and deletion of advertisements by recruiters",
         "Multi-language internationalization (i18n) of the interface"],
        ["Online application with PDF CV upload and application history",
         "Mobile native applications (iOS / Android)"],
        ["Administration back-office: recruiter validation, offer moderation",
         "Payment gateway for sponsored offers"],
        ["Advanced multi-criteria search with pagination",
         "Real-time chat between candidate and recruiter"],
        ["Dashboards and statistics for recruiters and administrators",
         "Machine-learning-based recommendation of offers"],
    ], widths=[Cm(8), Cm(8)])

h3(doc, "1.9. Stakeholder analysis")
add_body(doc,
    "The stakeholder analysis identifies the parties interested in the project, their expectations and their level of "
    "influence, in order to manage communication and priorities throughout the project.")
add_table_caption(doc, "Stakeholder analysis")
make_table(doc, ["Stakeholder", "Role / interest", "Expectation", "Influence"],
    [
        ["Candidate (end user)", "Browse and apply to offers", "Smooth, fast, secure application", "High"],
        ["Recruiter (end user)", "Publish and track offers", "Efficient advertisement management", "High"],
        ["Administrator", "Moderate and steer", "Full control of content", "High"],
        ["Academic supervisor", "Guide and evaluate", "Rigor and methodology", "High"],
        ["Jury members", "Evaluate the work", "Clear, complete report", "Medium"],
        ["Hosting provider (IIS)", "Host the application", "Stable deployment", "Low"],
    ], widths=[Cm(3.5), Cm(4.0), Cm(5.0), Cm(2.0)])

h3(doc, "1.10. SMART objectives")
add_body(doc,
    "The objectives of the project were formulated according to the SMART criteria (Specific, Measurable, Achievable, "
    "Relevant, Time-bound) in order to guarantee their clarity and their evaluability at the end of the project.")
add_table_caption(doc, "SMART objectives of the project")
make_table(doc, ["Objective", "SMART criterion", "Target"],
    [
        ["Deliver an operational web platform", "Specific & Achievable", "End of week 12"],
        ["Cover 7 functional modules", "Measurable", "7 sprints delivered"],
        ["Reach sub-3s response time on search", "Measurable", "P95 < 3s on 1000 offers"],
        ["Ensure role-based secure access", "Relevant", "Auth + roles implemented"],
        ["Produce the UML modeling", "Time-bound", "Diagrams per sprint"],
        ["Provide a moderation back-office", "Specific", "Admin module delivered"],
    ], widths=[Cm(5), Cm(5), Cm(5)])

h3(doc, "1.11. Risk identification")
add_body(doc,
    "Identifying risks from the start of the project makes it possible to put preventive actions in place. The following "
    "table summarizes the main identified risks, their impact, their probability and the planned mitigation measures. "
    "Each risk is given a severity score computed as impact times probability, which helps prioritize the monitoring.")
add_table_caption(doc, "Risk matrix")
make_table(doc, ["Risk", "Impact", "Probability", "Severity", "Preventive action"],
    [
        ["Schedule slippage", "High", "Medium", "6", "Breakdown into short sprints, weekly monitoring"],
        ["Poorly understood requirements", "High", "Low", "3", "Requirements validation with the supervisor at project start"],
        ["Technical complexity (ASP.NET MVC)", "Medium", "Medium", "4", "Technology watch, skill-building on ASP.NET MVC"],
        ["Data loss", "High", "Low", "3", "Regular backups, redeployment scripts"],
        ["Integration issues between modules", "Medium", "Medium", "4", "Integration testing per sprint"],
        ["Bugs in production", "High", "Low", "3", "Testing and validation phase before delivery"],
        ["File upload vulnerabilities", "High", "Medium", "6", "Type/size checks, restricted upload folder"],
        ["Performance on large offer volumes", "Medium", "Medium", "4", "Pagination (OFFSET/FETCH), indexed columns"],
        ["Scope creep", "Medium", "Medium", "4", "Strict backlog, change request validation"],
        ["Single-developer bottleneck", "Medium", "High", "6", "Time-boxing, prioritization, supervisor reviews"],
        ["Environment / IIS misconfiguration", "Medium", "Low", "2", "Deployment guide, local setup script"],
        ["Personal data leakage", "High", "Low", "3", "Hashed passwords, restricted access, no clear-text storage"],
    ], widths=[Cm(4.0), Cm(1.8), Cm(1.8), Cm(1.6), Cm(5.5)])

h3(doc, "1.12. Development methodologies")
add_body(doc,
    "Software development relies on different methodologies, generally classified into two broad families: classical "
    "methodologies (waterfall, V-model) and agile methodologies. Classical methodologies structure the project into "
    "sequential phases and require a complete specification of requirements from the start. They are suited to projects "
    "whose requirements are stable and well defined.")
add_body(doc,
    "Agile methodologies, on the contrary, advocate an iterative and incremental approach, based on close collaboration "
    "with the client, adaptation to change and frequent delivery of working software. Among agile frameworks, SCRUM is "
    "one of the most widespread. It breaks the project down into short iterations called sprints, framed by ceremonies "
    "(sprint planning, daily scrum, sprint review, retrospective) and well-defined roles (Product Owner, Scrum Master, "
    "development team).")
add_body(doc,
    "For our project, we chose the SCRUM methodology. This choice is justified by the nature of the project, whose "
    "requirements were progressively refined, and by the need to deliver regular, evaluable functional increments. The "
    "product backlog was prioritized and then split into seven sprints grouped into two releases.")
add_body(doc,
    "The life cycle adopted for each sprint is the following: analysis of the sprint requirements, design (UML diagrams), "
    "implementation (coding and interfaces), testing and validation, then conclusion and demonstration. This cycle repeats "
    "for each sprint and allows a progressive ramp-up of the platform.")

h3(doc, "1.13. Internal monitoring")
add_body(doc,
    "Project monitoring was ensured through weekly meetings with the supervisor, during which progress, difficulties and "
    "directions were discussed. A task tracking board (backlog) was kept up to date in order to visualize the progress "
    "of each feature. Source code versions were versioned with Git, and each sprint ended with a review of the produced "
    "increment and a short retrospective to identify improvement actions for the next iteration.")

h3(doc, "1.14. General quality requirements")
add_body(doc,
    "The quality plan translates the general quality requirements into measurable acceptance criteria. It was used as a "
    "reference during the validation of each sprint.")
add_table_caption(doc, "Quality plan")
make_table(doc, ["Quality criterion", "Acceptance criterion", "Verification means"],
    [
        ["Reliability", "No unhandled exception on main scenarios", "Manual + integration testing"],
        ["Security", "Passwords hashed, role-based access", "Code review, auth tests"],
        ["Ergonomics", "Responsive layout, intuitive navigation", "UI review on Chrome/Edge/Firefox"],
        ["Performance", "P95 response time < 3s on 1000 offers", "Search benchmark with sample data"],
        ["Maintainability", "Layered code (DAL/BLL/Controllers/Views)", "Architecture review"],
        ["Portability", "Compatibility with modern browsers", "Cross-browser checks"],
        ["Data integrity", "Referential integrity via foreign keys", "DB constraints + tests"],
    ], widths=[Cm(3.5), Cm(6.5), Cm(4.5)])
for q in [
    "Reliability: the application must run without errors on the main scenarios.",
    "Security: passwords must be stored securely, access restricted by roles.",
    "Ergonomics: clear interface, intuitive navigation, responsive design.",
    "Performance: acceptable response time even with a large volume of offers.",
    "Maintainability: code structured in layers (DAL, BLL, Controllers, Views), naming conventions.",
    "Portability: compatibility with modern browsers (Chrome, Edge, Firefox).",
]:
    bullet(doc, q)

h3(doc, "1.15. Project schedule")
add_body(doc,
    "The provisional schedule organizes the project over twelve weeks, from the requirements analysis to the final "
    "delivery. The detailed breakdown by phase and by sprint is given in the table below.")
add_table_caption(doc, "Project schedule")
make_table(doc, ["Phase", "Period", "Deliverable", "Owner"],
    [
        ["Requirements analysis", "Week 1-2", "Specifications document", "Developer + Supervisor"],
        ["Technical study", "Week 3", "Technology choices, architecture", "Developer"],
        ["Global design and modeling", "Week 3-4", "Global UML diagrams", "Developer"],
        ["Sprint 1 - Authentication", "Week 4", "Profile module", "Developer"],
        ["Sprint 2 - Offer management", "Week 5", "Offers module", "Developer"],
        ["Sprint 3 - Recruiter space", "Week 6", "Recruiter module", "Developer"],
        ["Sprint 4 - Administration", "Week 7", "Admin module", "Developer"],
        ["Release 1 - Review", "Week 7", "Release 1 demo", "Developer + Supervisor"],
        ["Sprint 5 - Applications", "Week 8", "Apply module", "Developer"],
        ["Sprint 6 - Search", "Week 9", "Search module", "Developer"],
        ["Sprint 7 - Statistics", "Week 10", "Dashboards", "Developer"],
        ["Release 2 - Review", "Week 10", "Release 2 demo", "Developer + Supervisor"],
        ["Testing and deployment", "Week 11-12", "Final delivery", "Developer"],
    ], widths=[Cm(4.0), Cm(2.5), Cm(4.5), Cm(3.5)])

h3(doc, "1.16. Actual project schedule")
add_body(doc,
    "The actual progress of the project globally followed the provisional schedule, with some adjustments on sprints 3 "
    "and 5 where the complexity of the business rules required a little extra time. These deviations were absorbed thanks "
    "to the flexibility of the agile approach and to the re-prioritization of certain backlog tasks.")

h3(doc, "1.17. Variance analysis")
add_body(doc,
    "The analysis of the variances between the provisional schedule and the actual progress reveals a minor shift of about "
    "5 to 7 days over the whole project, mainly attributable to the skill-building on certain technical aspects (file "
    "upload, pagination) and to the testing phase. No major feature was sacrificed; the content of the seven sprints was "
    "fully delivered.")

h3(doc, "1.18. Project organization")
add_body(doc,
    "The project was carried out solo, under the supervision of Mr. Mohamed-Amine HAMMA. The organization of the work "
    "relied on the following tools: Visual Studio as the development environment, SQL Server Management Studio for the "
    "database, Git for versioning, and Enterprise Architect for UML modeling. Communication with the supervisor was "
    "weekly, complemented by written progress notes.")
add_body(doc, "Conclusion")
add_body(doc,
    "This chapter was devoted to presenting the general context of the JobzFactory project. We defined the problem "
    "statement, the objectives, the missions and the deliverables, benchmarked the existing solutions, carried out a SWOT "
    "analysis, framed the scope and the stakeholders, formulated SMART objectives, identified and prioritized the risks, "
    "justified the choice of the agile SCRUM methodology and described the organization, the quality plan and the schedule "
    "of the project. The following chapter is devoted to the detailed requirements analysis.")
doc.add_page_break()

# ===================== CHAPTER 2 =====================
chapter_title(doc, "CHAPTER 2: Requirements Analysis")
h2(doc, "1. General problem statement reminder")
add_body(doc,
    "The problem consists in designing a unified web-based recruitment platform offering a smooth experience to candidates "
    "and recruiters, ensuring rigorous application tracking and providing a back-office of administration for the moderation "
    "and steering of the content. This chapter formalizes this problem into a structured set of actors, use cases, "
    "functional and non-functional requirements and business rules.")

h2(doc, "2. Project objective reminder")
add_body(doc,
    "The objective is to deliver an operational web application allowing candidates to manage their profile and apply, "
    "recruiters to publish and track their advertisements, and administrators to moderate and steer the platform. The "
    "requirements analysis is the prerequisite for any reliable design and implementation work.")

h2(doc, "3. Description of the project business processes")
add_body(doc,
    "The main business process of JobzFactory is broken down into several flows. The candidate flow: registration, account "
    "activation, profile completion, offer search, detail browsing, application with CV upload. The recruiter flow: "
    "registration, validation by the administration, publication of an advertisement, tracking of received applications, "
    "update or deletion of the advertisement. The administration flow: validation of recruiters, moderation of offers, "
    "deletion of unwanted content, consultation of statistics. These three flows are interconnected: an offer published by "
    "a recruiter becomes visible to candidates only after validation by the administration, and an application submitted by "
    "a candidate is reflected in the recruiter's tracking view.")

h2(doc, "4. Actor identification")
h3(doc, "4.1. Actors presentation")
add_body(doc,
    "The system identifies four main actors. The Visitor is an unauthenticated user who can browse and search offers and "
    "start a registration. The Candidate is an authenticated user with a profile who can apply to offers. The Recruiter is "
    "an authenticated user, validated by the administration, who publishes and manages advertisements. The Administrator is "
    "a super-user who manages recruiters, moderates offers and consults statistics. A secondary actor, the Mail Server, "
    "intervenes for account activation emails.")
add_table_caption(doc, "Actor identification")
make_table(doc, ["Actor", "Type", "Description", "Main actions"],
    [
        ["Visitor", "Primary, unauthenticated", "Any internet user landing on the platform", "Browse, search, view offers, register"],
        ["Candidate", "Primary, authenticated", "Registered user with a profile and CV", "Manage profile, apply, view history"],
        ["Recruiter", "Primary, authenticated", "Validated company representative", "Publish / edit / delete offers, view applications"],
        ["Administrator", "Primary, authenticated", "Platform super-user", "Validate recruiters, moderate offers, view stats"],
        ["Mail Server", "Secondary, system", "Email sending service", "Send activation / confirmation emails"],
        ["Database (SQL Server)", "Secondary, system", "Persistence service", "Store and retrieve all business data"],
    ], widths=[Cm(2.8), Cm(3.0), Cm(5.0), Cm(4.2)])

h3(doc, "4.2. Actor roles")
bullet(doc, "Visitor: browse the home page, search and view offers, register.")
bullet(doc, "Candidate: all visitor actions, plus manage their profile, apply, view their application history.")
bullet(doc, "Recruiter: manage their account, publish/edit/delete advertisements, view received applications.")
bullet(doc, "Administrator: validate recruiters, moderate offers, manage recruiters, view statistics.")
bullet(doc, "Mail Server: deliver account activation and confirmation messages.")

h2(doc, "5. Use case identification")
add_body(doc,
    "The use cases of the platform are organized by module. The table below lists all the identified use cases, the actor "
    "that triggers them and the module to which they belong. This identification constitutes the backbone of the product "
    "backlog that is later split into sprints.")
add_table_caption(doc, "Use case identification by module")
make_table(doc, ["ID", "Use case", "Actor", "Module"],
    [
        ["UC1", "Register as a candidate", "Visitor", "Profile"],
        ["UC2", "Activate account", "Candidate", "Profile"],
        ["UC3", "Log in / Log out", "Candidate / Recruiter / Admin", "Profile"],
        ["UC4", "Manage profile and CV", "Candidate", "Profile"],
        ["UC5", "View the list of offers", "Visitor / Candidate", "Offers"],
        ["UC6", "View the detail of an offer", "Visitor / Candidate", "Offers"],
        ["UC7", "Search and filter offers", "Visitor / Candidate", "Search"],
        ["UC8", "Apply to an offer", "Candidate", "Applications"],
        ["UC9", "Upload a CV (PDF)", "Candidate", "Applications"],
        ["UC10", "View application history", "Candidate", "Applications"],
        ["UC11", "Register as a recruiter", "Visitor", "Recruiter"],
        ["UC12", "Publish an advertisement", "Recruiter", "Recruiter"],
        ["UC13", "Edit an advertisement", "Recruiter", "Recruiter"],
        ["UC14", "Delete an advertisement", "Recruiter", "Recruiter"],
        ["UC15", "View received applications", "Recruiter", "Recruiter"],
        ["UC16", "View advertisement history", "Recruiter", "Recruiter"],
        ["UC17", "Validate a recruiter", "Administrator", "Administration"],
        ["UC18", "Moderate an offer", "Administrator", "Administration"],
        ["UC19", "Delete a recruiter / content", "Administrator", "Administration"],
        ["UC20", "View dashboards and statistics", "Administrator / Recruiter", "Statistics"],
    ], widths=[Cm(1.5), Cm(5.5), Cm(4.0), Cm(3.0)])

h2(doc, "6. Functional requirements")
add_body(doc,
    "The functional requirements describe what the system must do. Each requirement is given an identifier (RF-xx), a "
    "priority (High / Medium / Low) and the module concerned. They are derived from the use cases listed above.")
add_table_caption(doc, "Functional requirements (RF)")
make_table(doc, ["ID", "Requirement", "Module", "Priority"],
    [
        ["RF-01", "Allow candidate registration with email uniqueness check", "Profile", "High"],
        ["RF-02", "Send an activation link by email and activate the account", "Profile", "High"],
        ["RF-03", "Provide authentication and session management", "Profile", "High"],
        ["RF-04", "Allow profile creation, editing and CV management", "Profile", "High"],
        ["RF-05", "Display the home page with the recent offers", "Offers", "High"],
        ["RF-06", "Display a paginated list of validated offers", "Offers", "High"],
        ["RF-07", "Display the detail of an offer with the recruiter info", "Offers", "High"],
        ["RF-08", "Provide a multi-criteria search (keyword, sector, city, country)", "Search", "High"],
        ["RF-09", "Paginate the search results", "Search", "Medium"],
        ["RF-10", "Allow a candidate to apply to an offer once", "Applications", "High"],
        ["RF-11", "Allow PDF CV upload with type and size checks", "Applications", "High"],
        ["RF-12", "Display the candidate's application history", "Applications", "Medium"],
        ["RF-13", "Allow recruiter registration and validation workflow", "Recruiter", "High"],
        ["RF-14", "Allow CRUD operations on advertisements", "Recruiter", "High"],
        ["RF-15", "Display the applications received per offer", "Recruiter", "High"],
        ["RF-16", "Allow the administrator to validate / reject recruiters", "Administration", "High"],
        ["RF-17", "Allow the administrator to moderate / delete offers", "Administration", "High"],
        ["RF-18", "Compute and display dashboards and statistics", "Statistics", "Medium"],
    ], widths=[Cm(1.5), Cm(8.0), Cm(3.0), Cm(2.0)])

h2(doc, "7. Non-functional requirements")
add_body(doc,
    "The non-functional requirements describe how the system must behave. They cover performance, security, usability and "
    "maintainability. They are used as acceptance criteria during the validation of each sprint.")
add_table_caption(doc, "Non-functional requirements (NFR)")
make_table(doc, ["ID", "Category", "Requirement", "Target"],
    [
        ["NFR-01", "Performance", "Search response time", "P95 < 3s on 1000 offers"],
        ["NFR-02", "Performance", "Page rendering", "First paint < 1.5s"],
        ["NFR-03", "Security", "Password storage", "Hashed (no clear text)"],
        ["NFR-04", "Security", "Access control", "Role-based, server-side enforced"],
        ["NFR-05", "Security", "File upload", "PDF only, size < 5 MB, sandboxed folder"],
        ["NFR-06", "Usability", "Responsive design", "Desktop + tablet, Bootstrap grid"],
        ["NFR-07", "Usability", "Navigation", "Consistent menu, breadcrumbs"],
        ["NFR-08", "Maintainability", "Layered architecture", "DAL / BLL / Controllers / Views"],
        ["NFR-09", "Maintainability", "Naming conventions", "PascalCase C#, snake_case SQL"],
        ["NFR-10", "Portability", "Browser compatibility", "Chrome, Edge, Firefox"],
    ], widths=[Cm(1.5), Cm(2.6), Cm(6.5), Cm(4.4)])

h2(doc, "8. Business rules")
add_body(doc,
    "The business rules constrain the behavior of the system and guarantee the coherence of the data and the processes. "
    "They are enforced both in the application layer and in the database constraints.")
add_table_caption(doc, "Business rules")
make_table(doc, ["ID", "Rule"],
    [
        ["BR-01", "A recruiter can publish an offer only after validation by the administration."],
        ["BR-02", "An offer must include a title, a description, a sector, a city and a country."],
        ["BR-03", "A candidate can apply only once to the same offer."],
        ["BR-04", "Uploading a CV is mandatory to apply; only PDF files are accepted (max 5 MB)."],
        ["BR-05", "The administrator can disable or delete any non-compliant offer."],
        ["BR-06", "Passwords are stored in hashed form and are never displayed in clear text."],
        ["BR-07", "Only offers with status 'validated' are publicly visible on the front-office."],
        ["BR-08", "Offers are sorted by publication date descending by default."],
        ["BR-09", "An account that is not activated within a reasonable delay remains in 'pending' state."],
        ["BR-10", "Each moderation action (validation / rejection / deletion) is logged with its date and author."],
        ["BR-11", "The candidate's email address is unique across the platform."],
        ["BR-12", "A recruiter cannot delete an offer that already received applications without confirmation."],
    ], widths=[Cm(1.8), Cm(13.2)])

h2(doc, "9. Global use case diagram")
add_body(doc,
    "The global use case diagram gives a synthetic view of all the use cases of the platform and of the relationships "
    "between the actors and these use cases. It is the entry point of the modeling and is detailed per sprint in the "
    "implementation chapters.")
add_figure(doc, "Global use case diagram of the JobzFactory platform",
           "[ UML use case diagram: Visitor, Candidate, Recruiter, Administrator and the 20 use cases ]",
           box_h=Cm(8))
add_body(doc, "Conclusion")
add_body(doc,
    "This chapter made it possible to formalize the requirements of the project by identifying the actors, the use cases, "
    "the functional and non-functional requirements and the business rules. The global use case diagram gives a synthetic "
    "view of the expected behavior of the platform. The following chapter deals with the technical study and the choice "
    "of development tools.")
doc.add_page_break()

# ===================== CHAPTER 3 =====================
chapter_title(doc, "CHAPTER 3: Technical Study")
h2(doc, "1. Technical analysis")
h3(doc, "1.1. Architecture patterns")
add_body(doc,
    "Several architectural patterns were considered before choosing the one adopted for JobzFactory. The Model-View-"
    "Controller (MVC) pattern separates the application into three components: the model, which carries the data and the "
    "business logic; the view, which renders the user interface; and the controller, which handles the incoming requests, "
    "orchestrates the model and selects the view to return. This separation favors a clean separation of concerns and "
    "improves testability.")
add_body(doc,
    "A layered architecture organizes the code into horizontal layers (presentation, business logic, data access) where "
    "each layer depends only on the layer immediately below it. The three-tier architecture is a deployment-oriented "
    "variant of the layered approach, where the presentation, the business and the data tiers can be deployed on separate "
    "physical tiers. For JobzFactory we adopt a layered architecture deployed on a single tier, with the ASP.NET MVC "
    "pattern on the presentation side.")
add_body(doc,
    "Compared with a monolithic WebForms-style approach, the MVC pattern offers a finer control over the generated HTML, "
    "a cleaner routing and a better support for automated testing. Compared with a client-side single-page-application "
    "approach, a server-rendered MVC application is simpler to deploy in an academic IIS environment and well aligned "
    "with the skills targeted by the Software Engineering Master's program.")

h3(doc, "1.2. Solution architecture")
add_body(doc,
    "JobzFactory is developed with the ASP.NET MVC framework, which enforces a clear separation of responsibilities "
    "between the model (data and business logic), the view (user interface) and the controller (request orchestration). "
    "This architecture promotes the maintainability, testability and evolvability of the application.")
add_body(doc,
    "The solution is organized into several Visual Studio projects corresponding to the different spaces of the "
    "application: JobzFactory (public candidate front-office), Recruteur (recruiter space), Administration (back-office) "
    "and Profil (candidate profile management). A data access layer JF.DAL mutualizes access to the SQL Server database. "
    "Each project has its own controllers, views and routing configuration, but they all rely on the shared JF.DAL layer "
    "and on the same SQL Server database.")
add_figure(doc, "Overall architecture of the JobzFactory solution",
           "[ Architecture diagram: Browser -> IIS / ASP.NET MVC (Controllers) -> BLL -> JF.DAL -> SQL Server ]",
           box_h=Cm(7))

h3(doc, "1.3. .NET Framework and the CLR")
add_body(doc,
    "The .NET Framework is a Microsoft software platform that provides a managed execution environment, a comprehensive "
    "class library (FCL) and support for multiple programming languages. The Common Language Runtime (CLR) is the "
    "execution engine of .NET: it loads the Intermediate Language (IL) produced by the compilers, compiles it to native "
    "code through the Just-In-Time (JIT) compiler, manages memory through garbage collection, enforces type safety and "
    "handles exceptions. This managed execution model guarantees a robust and secure runtime for the JobzFactory "
    "application.")
add_body(doc,
    "C# source files are compiled into assemblies containing IL metadata. At runtime the CLR resolves references between "
    "the assemblies of the solution (JobzFactory, Recruteur, Administration, Profil, JF.DAL) and ensures that the types "
    "exposed by the data access layer are correctly consumed by the controllers of each presentation project.")

h3(doc, "1.4. ADO.NET data access")
add_body(doc,
    "Data access is implemented with ADO.NET, the data access technology of the .NET Framework. ADO.NET relies on "
    "connected objects (SqlConnection, SqlCommand, SqlDataReader) for streaming read access and on disconnected objects "
    "(DataSet, DataTable, SqlDataAdapter) for offline manipulation of data. In JobzFactory, the JF.DAL layer centralizes "
    "the connection string management, the opening and closing of connections and the execution of parameterized queries "
    "and stored procedures. The use of parameterized queries prevents SQL injection and improves the readability and the "
    "maintainability of the data access code.")

h3(doc, "1.5. IIS hosting")
add_body(doc,
    "Internet Information Services (IIS) is the Microsoft web server used to host the ASP.NET MVC application. IIS "
    "manages the application pools, the bindings, the request pipeline and the static and dynamic content. Each "
    "presentation project (JobzFactory, Recruteur, Administration, Profil) is published as a separate IIS web site, "
    "sharing the same application pool version and the same SQL Server database. This deployment topology makes it "
    "possible to expose distinct entry points (public front-office, recruiter space, administration back-office) while "
    "mutualizing the data and the business layer.")

h3(doc, "1.6. Razor views")
add_body(doc,
    "Razor is the view engine used by ASP.NET MVC. It allows mixing HTML markup with C# code using a compact @-syntax "
    "inside .cshtml files. Razor views receive a strongly-typed model from the controller and produce the HTML response. "
    "Layouts and partial views are used to mutualize the structural elements (header, navigation, footer) across the "
    "pages of each project, ensuring a consistent look and feel and reducing duplication.")

h3(doc, "1.7. Authentication and session mechanism")
add_body(doc,
    "Authentication is implemented with ASP.NET forms authentication combined with session state. When a user logs in, "
    "the controller checks the credentials against the database, sets the authentication cookie and stores the user "
    "context in the session. Each request to a protected area is intercepted by an authorization filter that checks the "
    "role stored in the session and redirects unauthenticated or unauthorized users to the login page or to an access "
    "denied page. Passwords are stored hashed and are never compared in clear text.")

h3(doc, "1.8. Deployment topology")
add_body(doc,
    "The deployment topology describes how the components of the solution are physically distributed. In our academic "
    "configuration, all the components run on a single Windows server hosting IIS and SQL Server. The diagram below "
    "illustrates the deployed deployment: the client browser communicates over HTTP with IIS, which hosts the four "
    "ASP.NET MVC sites, all relying on the JF.DAL layer and on the SQL Server database; the mail server is reached over "
    "SMTP for the activation emails.")
add_figure(doc, "Deployed deployment diagram of the JobzFactory platform",
           "[ UML deployment diagram: Browser -> IIS {JobzFactory, Recruteur, Administration, Profil} -> JF.DAL -> SQL Server ; SMTP -> Mail Server ]",
           box_h=Cm(8))

h3(doc, "1.9. The tools used")
h3brown(doc, "1.9.1. ASP.NET MVC framework")
add_body(doc,
    "ASP.NET MVC is a Microsoft web framework based on the Model-View-Controller pattern. It offers fine control over "
    "the generated HTML, flexible URL routing and an organization of the code into controllers and Razor views (.cshtml). "
    "It integrates model binding, validation through data annotations and filter-based cross-cutting concerns "
    "(authorization, exception handling, logging).")
h3brown(doc, "1.9.2. C# language")
add_body(doc,
    "C# is an object-oriented, strongly-typed, safe and productive programming language, running on the .NET platform. "
    "It was used for all the business logic and the controllers. Its features (generics, LINQ, async/await, attributes) "
    "make it possible to write concise and robust code for the data access and the orchestration layers.")
h3brown(doc, "1.9.3. SQL Server")
add_body(doc,
    "Microsoft SQL Server is the chosen relational database management system. It ensures the persistent storage of "
    "offers, profiles, recruiters, applications and references (sectors, cities, countries). Data access is performed "
    "via ADO.NET within the JF.DAL layer. Referential integrity is enforced through foreign keys and constraints.")
add_figure(doc, "SQL Server Management Studio logo", "[ SSMS logo ]", box_h=Cm(6))
h3brown(doc, "1.9.4. Web tools used")
add_body(doc,
    "The user interface relies on HTML5 for the structure, CSS3 for the styling, JavaScript for the interactivity and "
    "Bootstrap for a responsive grid system and pre-styled components. jQuery is used for DOM manipulation and AJAX "
    "calls where needed.")
add_figure(doc, "Logos of the web technologies used (HTML5, CSS3, JS, Bootstrap)",
           "[ Logos HTML5, CSS3, JavaScript, Bootstrap, jQuery ]", box_h=Cm(6))
h3brown(doc, "1.9.5. Visual Studio")
add_body(doc,
    "Microsoft Visual Studio is the integrated development environment used. It offers a powerful code editor, debugging, "
    "project management, NuGet package management and publishing to IIS. It was used to structure the multi-project "
    "solution and to run the application locally during development.")
add_figure(doc, "Microsoft Visual Studio logo", "[ Visual Studio logo ]", box_h=Cm(6))
h3brown(doc, "1.9.6. Enterprise Architect")
add_body(doc,
    "Enterprise Architect from Sparx Systems is the UML modeling tool used to produce the use case, class, sequence and "
    "deployment diagrams of the project. It supports the standard UML 2.x notation and makes it possible to keep a "
    "coherent model across the sprints.")
add_figure(doc, "Enterprise Architect logo", "[ Enterprise Architect logo ]", box_h=Cm(6))
h3brown(doc, "1.9.7. Git")
add_body(doc,
    "Git was used as a distributed version control system, in order to track the evolution of the source code, keep a "
    "history of modifications and make it easier to roll back when needed. Commits were organized by feature and by "
    "sprint, and branches were used to isolate experimental work from the main line.")
add_figure(doc, "Git logo", "[ Git logo ]", box_h=Cm(6))
h3brown(doc, "1.9.8. SQL Server Management Studio")
add_body(doc,
    "SSMS is the database administration tool. It was used to create the schema, enter the reference data, run the "
    "deployment scripts and inspect the queries' execution plans in order to validate the indexing and pagination "
    "strategy.")
add_body(doc, "Conclusion")
add_body(doc,
    "This chapter was devoted to the technical study of the project. We justified the architecture choices by comparing "
    "the MVC, layered and three-tier patterns, presented the .NET execution model (CLR, IL, JIT), the ADO.NET data "
    "access, the IIS hosting, the Razor views, the authentication and session mechanism and the deployment topology. "
    "We also presented all the tools used for development, data storage, modeling and versioning. The following chapter "
    "presents the global design and modeling of the platform.")
doc.add_page_break()

# ===================== CHAPTER 4: GLOBAL DESIGN =====================
chapter_title(doc, "CHAPTER 4: Global Design and Modeling")
add_body(doc,
    "Before detailing the implementation sprint by sprint, this chapter presents the global design and modeling of the "
    "JobzFactory platform. It gathers the cross-cutting UML diagrams that give a synthetic view of the whole system: the "
    "global class diagram, the global use case diagram and the package / architecture diagram. It also describes the main "
    "business entities and the relationships between them.")

h2(doc, "1. Global class diagram")
add_body(doc,
    "The global class diagram describes the static structure of the information system. It groups the main business "
    "classes of the platform and the associations between them. The classes are organized around the candidate, the "
    "recruiter, the offer and the application, complemented by the reference classes (ActivitySector, City, Country) and "
    "the administration classes (Administrator, OfferAction).")
add_figure(doc, "Global class diagram of the JobzFactory platform",
           "[ UML class diagram: User, Candidate, Recruiter, Administrator, Profile, Offer, Application, CVFile, ActivitySector, City, Country, OfferAction ]",
           box_h=Cm(8))
add_body(doc,
    "The User class is the root of the authentication model. Candidate and Recruiter specialize the user according to "
    "the role. The Profile and CVFile classes are attached to the candidate. The Offer class is published by a Recruiter "
    "and references one ActivitySector, one City and one Country. The Application class links a Candidate to an Offer and "
    "carries the uploaded CVFile. The OfferAction class logs the moderation actions performed by the Administrator.")

h2(doc, "2. Global use case diagram")
add_body(doc,
    "The global use case diagram, already introduced in the requirements chapter, is recalled here from the design "
    "perspective. It organizes the twenty use cases around the four actors and highlights the 'include' and 'extend' "
    "relationships: applying to an offer includes uploading a CV; moderating an offer extends the validation use case.")
add_figure(doc, "Global use case diagram (design view)",
           "[ UML use case diagram with include/extend relationships between the 20 use cases ]",
           box_h=Cm(8))

h2(doc, "3. Package and architecture diagram")
add_body(doc,
    "The package diagram shows the organization of the solution into projects and packages. The presentation packages "
    "(JobzFactory, Recruteur, Administration, Profil) depend on the shared JF.DAL package, which itself depends on the "
    "System.Data.SqlClient primitives and on the database. Cross-cutting packages (Models, ViewModels, Helpers) are "
    "referenced where needed.")
add_figure(doc, "Package / architecture diagram of the solution",
           "[ UML package diagram: JobzFactory, Recruteur, Administration, Profil -> JF.DAL -> System.Data / SQL Server ]",
           box_h=Cm(8))

h2(doc, "4. Main entities and relationships")
add_body(doc,
    "The main entities of the platform and their relationships are described below. A candidate owns one profile and may "
    "own one or several CV files; a recruiter publishes zero to many offers; an offer belongs to one sector, one city and "
    "one country and may receive zero to many applications; an application is submitted by one candidate for one offer and "
    "references one uploaded CV file. The administrator performs moderation actions on offers and validation actions on "
    "recruiters; each action is logged in the OfferAction table with its type, its date and its author.")
add_body(doc,
    "The cardinalities are the following. User (1) -- (0..1) Profile; Candidate (1) -- (0..*) CVFile; Recruiter (1) -- "
    "(0..*) Offer; Offer (1) -- (0..*) Application; Candidate (1) -- (0..*) Application; Application (1) -- (1) CVFile "
    "for the considered application; ActivitySector (1) -- (0..*) Offer; City (1) -- (0..*) Offer; Country (1) -- (0..*) "
    "City; Administrator (1) -- (0..*) OfferAction; Offer (1) -- (0..*) OfferAction. These cardinalities are translated "
    "into primary and foreign keys in the physical database model described in the sprint data dictionaries and in the "
    "appendix.")
add_body(doc,
    "The design also defines enumerations that constrain the value of certain attributes: the account state (Pending, "
    "Active, Disabled), the offer status (Pending, Validated, Rejected, Deleted) and the moderation action type "
    "(Validation, Rejection, Deletion). These enumerations are implemented both as check constraints at the database "
    "level and as constants in the application layer.")
add_body(doc, "Conclusion")
add_body(doc,
    "This chapter presented the global design and modeling of the platform through the global class diagram, the global "
    "use case diagram and the package / architecture diagram. It also described the main entities and the cardinalities "
    "of their relationships. This global view provides the framework in which the sprints described in the following "
    "chapters are progressively implemented.")
doc.add_page_break()

# ===================== SPRINT HELPER =====================
def sprint(doc, title, module_paras, fiche_rows, backlog_rows, mgmt_rules,
           actor_rows, uc_spec_rows, uc_diagram_desc, class_desc,
           seq1_title, seq1_desc, seq2_title, seq2_desc,
           dd_rows, iface_paras, iface1_desc, iface2_desc, conclusion_text):
    """Render a full, expanded sprint section (~5-7 pages)."""
    h1(doc, title)
    hr(doc)

    # 1. Sprint description
    h2(doc, "1. Sprint description")
    h3(doc, "1.1. Module overview")
    for para in module_paras:
        add_body(doc, para)
    h3(doc, "1.2. Sprint sheet")
    add_table_caption(doc, "Sprint sheet")
    make_table(doc, ["Item", "Value"], fiche_rows, widths=[Cm(5), Cm(10)])
    h3(doc, "1.3. Sprint backlog")
    add_table_caption(doc, "Sprint backlog")
    make_table(doc, ["ID", "User Story", "Priority", "Estimate", "Status"],
               backlog_rows, widths=[Cm(1.5), Cm(7.0), Cm(2.0), Cm(1.8), Cm(2.2)])
    h3(doc, "1.4. Management rules")
    add_body(doc,
        "The following management rules structure the running of the sprint. They frame the priorities, the definition of "
        "done and the working agreements adopted during the iteration.")
    for r in mgmt_rules:
        numbered(doc, r)
    h3(doc, "1.5. Sprint actors")
    add_body(doc,
        "The actors involved in this sprint are listed in the table below, with their role and their main interaction "
        "with the module.")
    add_table_caption(doc, "Sprint actors")
    make_table(doc, ["Actor", "Role", "Main interaction"], actor_rows,
               widths=[Cm(3.0), Cm(5.0), Cm(7.0)])

    # 2. Analysis and design
    h2(doc, "2. Analysis and design")
    h3(doc, "2.1. Use case diagram")
    add_body(doc,
        "The use case diagram of the sprint scopes the functionalities to be delivered during the iteration. It identifies "
        "the actors concerned and the use cases that they trigger.")
    add_figure(doc, "Sprint use case diagram", uc_diagram_desc, box_h=Cm(7))
    h3(doc, "2.2. Use case description")
    add_body(doc,
        "The main use case of the sprint is described below in a structured specification table. The table details the "
        "actor, the preconditions, the main flow, the alternative flows and the postcondition.")
    add_table_caption(doc, "Use case specification")
    make_table(doc, ["Item", "Description"], uc_spec_rows,
               widths=[Cm(3.5), Cm(11.5)], font_size=10)
    h3(doc, "2.3. Class diagram")
    add_body(doc,
        "The class diagram of the sprint details the static structure of the classes involved in the module and the "
        "associations between them.")
    add_figure(doc, "Sprint class diagram", class_desc, box_h=Cm(7))
    h3(doc, "2.4. Sequence diagram (use case 1)")
    add_body(doc, seq1_title)
    add_figure(doc, "Sprint sequence diagram (use case 1)",
               "[ UML sequence diagram: " + seq1_desc + " ]", box_h=Cm(8))
    add_body(doc, seq1_desc)
    h3(doc, "2.5. Sequence diagram (use case 2)")
    add_body(doc, seq2_title)
    add_figure(doc, "Sprint sequence diagram (use case 2)",
               "[ UML sequence diagram: " + seq2_desc + " ]", box_h=Cm(8))
    add_body(doc, seq2_desc)
    h3(doc, "2.6. Data dictionary")
    add_body(doc,
        "The data dictionary lists the columns of the main tables manipulated by the sprint, with their type and a short "
        "description. It complements the class diagram from the physical storage perspective.")
    add_table_caption(doc, "Sprint data dictionary")
    make_table(doc, ["Field", "Type", "Description"], dd_rows,
               widths=[Cm(4.0), Cm(3.5), Cm(7.5)])

    # 3. Implementation
    h2(doc, "3. Implementation")
    h3(doc, "3.1. Description of the interfaces")
    for para in iface_paras:
        add_body(doc, para)
    add_figure(doc, "Sprint interfaces (list / form view)", iface1_desc, box_h=Cm(8))
    add_figure(doc, "Sprint interfaces (detail / dashboard view)", iface2_desc, box_h=Cm(8))

    # 4. Conclusion
    h2(doc, "4. Conclusion")
    add_body(doc, conclusion_text)
    doc.add_page_break()


# ===================== CHAPTER 5: RELEASE 1 =====================
chapter_title(doc, "CHAPTER 5: « Release 1 » (Sprints 1 to 4)")
add_body(doc,
    "The first release groups the four sprints that lay the foundations of the platform: authentication and candidate "
    "profile management, public offer management, the recruiter space and the administration back-office. At the end of "
    "this release, a complete vertical slice of the platform is operational, from the candidate registration to the "
    "moderation of the published content.")

# ---- Sprint 1
sprint(doc,
    "Sprint 1: Authentication and candidate profile management",
    [
        "This sprint covers registration, authentication and candidate profile management. It constitutes the entry point "
        "of the platform and lays the foundations of security and account management. The deliverable of this sprint is "
        "a fully functional authentication module allowing a visitor to become a candidate with an activated account.",
        "From an architectural point of view, this sprint introduces the Profil project, the User and Profile entities "
        "and the JF.DAL data access methods dedicated to authentication. It also sets up the session and authorization "
        "mechanisms that will be reused by all the protected areas of the platform.",
        "The validation of the sprint is based on the demonstration of the registration, activation, login and profile "
        "completion scenarios, and on the verification that an unauthenticated user cannot access the protected pages.",
    ],
    [("Sprint", "Sprint 1"), ("Release", "Release 1"), ("Duration", "1 week"),
     ("Objective", "Authenticate and manage the candidate profile"),
     ("Actors", "Visitor, Candidate, Administrator"),
     ("Entry criteria", "Requirements validated, environment ready"),
     ("Exit criteria", "Auth scenarios demonstrated and tested")],
    [("US1.1", "As a visitor, I register as a candidate", "High", "3", "Done"),
     ("US1.2", "As a candidate, I activate my account by email", "High", "2", "Done"),
     ("US1.3", "As a candidate, I log in", "High", "2", "Done"),
     ("US1.4", "As a candidate, I log out", "Medium", "1", "Done"),
     ("US1.5", "As a candidate, I fill in my profile", "Medium", "3", "Done"),
     ("US1.6", "As a candidate, I change my password", "Medium", "2", "Done"),
     ("US1.7", "As a system, I prevent access to protected pages for guests", "High", "2", "Done")],
    ["Each sprint starts with a planning meeting and ends with a review and a retrospective.",
     "A user story is considered done only when it is coded, tested and demonstrated.",
     "The backlog is frozen at the end of the sprint planning.",
     "Any new requirement is added to the product backlog, not to the running sprint.",
     "Credentials are never logged or displayed in clear text.",
     "The supervisor is informed at the end of each sprint of the produced increment."],
    [["Visitor", "Unauthenticated user", "Fills in the registration form"],
     ["Candidate", "Authenticated user", "Activates the account, logs in, manages the profile"],
     ["Mail Server", "Secondary actor", "Sends the activation email"],
     ["Administrator", "Validator", "May disable an account if needed"]],
    [["Use case name", "Register as a candidate"],
     ["Actor", "Visitor"],
     ["Precondition", "The visitor is not yet registered and the email is not used"],
     ["Main flow",
      "1. The visitor opens the registration page.\n2. The visitor fills in the form (email, password, name).\n3. The controller validates the inputs.\n4. The DAL creates the account in 'pending' state.\n5. An activation email is sent.\n6. The view displays a confirmation message."],
     ["Alternative flows",
      "A1: the email is already used -> an error message is displayed.\nA2: a field is invalid -> the form is redisplayed with the errors."],
     ["Postcondition", "The account exists in 'pending' state and an activation email has been sent"]],
    "[ UML use case diagram: Register, Activate account, Log in, Log out, Manage profile ]",
    "[ UML class diagram: classes User, Candidate, Profile, Session, ActivationToken ]",
    "'Sign up' sequence diagram",
    "The visitor fills in the registration form, the controller checks the uniqueness of the email, the service creates "
    "the account, generates an activation token and sends an email; the candidate confirms via the activation link.",
    "'Log in' sequence diagram",
    "The candidate enters their email and password, the controller retrieves the account, verifies the hashed password "
    "and the 'active' state, sets the authentication cookie and the session, then redirects to the profile space.",
    [("Email", "varchar(100)", "Unique identifier of the candidate"),
     ("Password", "varchar(255)", "Hashed password"),
     ("LastName", "varchar(50)", "Candidate last name"),
     ("FirstName", "varchar(50)", "Candidate first name"),
     ("Phone", "varchar(20)", "Phone number"),
     ("RegistrationDate", "datetime", "Account creation date"),
     ("AccountState", "varchar(20)", "State: pending / active / disabled"),
     ("ActivationToken", "varchar(100)", "Unique activation token"),
     ("TokenExpiry", "datetime", "Activation token expiry date"),
     ("LastLoginDate", "datetime", "Last successful login date")],
    ["The implemented interfaces include the registration form, the activation page, the login page and the profile "
     "completion form. An activation email sending feature was integrated. The forms are validated both client-side and "
     "server-side, and the error messages are clearly displayed next to the concerned fields.",
     "The profile space gives access to the personal information edition and to the password change form. A logout link "
     "is available in the navigation bar of all the protected pages."],
    "[ Screenshot: registration form and login page ]",
    "[ Screenshot: profile completion and password change views ]",
    "This sprint established the authentication foundation and candidate profile management. The registration, "
    "activation, login and profile management features are operational and secure. The session and authorization "
    "mechanisms introduced here are reused by all the following sprints.")

# ---- Sprint 2
sprint(doc,
    "Sprint 2: Offer management (public front-office)",
    [
        "This sprint covers the display, browsing and search of job offers on the public part of the platform, accessible "
        "to visitors and candidates. It is the showcase of the platform and the main entry point for a candidate looking "
        "for an offer.",
        "The sprint introduces the Offer, ActivitySector, City and Country entities and the read methods of the JF.DAL "
        "layer dedicated to the public display. Only offers with the 'validated' status are visible, which guarantees the "
        "quality of the published content.",
        "The validation of the sprint is based on the demonstration of the home page, the paginated offer list and the "
        "offer detail page, with the references (sector, city, country) and the issuing recruiter correctly displayed.",
    ],
    [("Sprint", "Sprint 2"), ("Release", "Release 1"), ("Duration", "1 week"),
     ("Objective", "Display and browse offers"),
     ("Actors", "Visitor, Candidate"),
     ("Entry criteria", "Reference data (sector, city, country) initialized"),
     ("Exit criteria", "Offer list and detail demonstrated")],
    [("US2.1", "As a visitor, I view the home page with recent offers", "High", "2", "Done"),
     ("US2.2", "As a visitor, I view the paginated list of offers", "High", "3", "Done"),
     ("US2.3", "As a visitor, I view the detail of an offer", "High", "2", "Done"),
     ("US2.4", "As a visitor, I search an offer by keyword", "Medium", "3", "Done"),
     ("US2.5", "As a visitor, I filter by sector/city/country", "Medium", "3", "Done"),
     ("US2.6", "As a visitor, I paginate the results", "Medium", "2", "Done")],
    ["Only validated offers are publicly visible.",
     "Offers are sorted by publication date descending by default.",
     "The number of offers per page is configurable.",
     "An empty result set displays a clear 'no offer' message.",
     "The home page highlights the most recent offers.",
     "The detail page must always show the recruiter information."],
    [["Visitor", "Unauthenticated user", "Browses and searches the offers"],
     ["Candidate", "Authenticated user", "Same as visitor, plus can apply later"],
     ["Database", "Persistence service", "Returns the validated offers"]],
    [["Use case name", "View the detail of an offer"],
     ["Actor", "Visitor / Candidate"],
     ["Precondition", "The offer exists and has the 'validated' status"],
     ["Main flow",
      "1. The visitor opens the offer list.\n2. The visitor clicks on an offer.\n3. The controller retrieves the offer by id.\n4. The DAL loads the references and the recruiter.\n5. The view displays the offer detail."],
     ["Alternative flows", "A1: the offer does not exist -> a 'not found' page is displayed."],
     ["Postcondition", "The offer detail is displayed with its references and recruiter"]],
    "[ UML use case diagram: View offers, View detail, Search, Filter, Paginate ]",
    "[ UML class diagram: classes Offer, ActivitySector, City, Country, Recruiter ]",
    "'View an offer' sequence diagram",
    "The visitor opens the home page, the controller retrieves the list of validated offers via the DAL, the view "
    "displays the paginated list; on click on an offer, the detail is loaded with its references and the issuing recruiter.",
    "'Paginate the offer list' sequence diagram",
    "The visitor clicks on a page number, the controller forwards the page index to the DAL which executes a paged "
    "query (OFFSET/FETCH), the view redisplays the list with the new page and the pagination controls.",
    [("OfferId", "int", "Offer identifier (PK)"),
     ("Title", "varchar(150)", "Job title"),
     ("Description", "text", "Detailed description"),
     ("SectorId", "int", "Activity sector (FK)"),
     ("CityId", "int", "Job city (FK)"),
     ("CountryId", "int", "Job country (FK)"),
     ("RecruiterId", "int", "Issuing recruiter (FK)"),
     ("PublicationDate", "datetime", "Posting date"),
     ("Status", "varchar(20)", "Offer status (validated/pending/rejected)"),
     ("ContractType", "varchar(30)", "Type of contract (CDI/CDD/...)"),
     ("SalaryRange", "varchar(50)", "Indicative salary range"),
     ("ViewsCount", "int", "Number of detail views")],
    ["The implemented interfaces include the modern home page with the list of recent offers, the offer listing page "
     "with pagination, and the offer detail page with information about the recruiter. The list page provides a search "
     "panel and filter dropdowns.",
     "The detail page presents the full description, the references, the recruiter information and an 'Apply' button "
     "that will be activated in the application sprint. The layout is responsive and adapts to tablet screens."],
    "[ Screenshot: home page and paginated offer list ]",
    "[ Screenshot: offer detail page with recruiter info ]",
    "This sprint provided visitors and candidates with a smooth navigation over offers. Display, filtering, pagination "
    "and detail browsing are operational and form the public showcase of the platform.")

# ---- Sprint 3
sprint(doc,
    "Sprint 3: Recruiter space",
    [
        "This sprint covers the space dedicated to recruiters: authentication, dashboard, publication, editing and deletion "
        "of advertisements, as well as browsing the received applications. It is the operational workspace of the "
        "recruiter and the main source of the content displayed on the public front-office.",
        "The sprint introduces the Recruiter, Advertisement and Application entities and the CRUD methods of the JF.DAL "
        "layer dedicated to the recruiter space. A recruiter can publish only after being validated by the administration, "
        "which guarantees a minimal control over the publishers.",
        "The validation of the sprint is based on the demonstration of the publication, edition and deletion of an "
        "advertisement, and on the browsing of the applications received per advertisement.",
    ],
    [("Sprint", "Sprint 3"), ("Release", "Release 1"), ("Duration", "1 week"),
     ("Objective", "Manage advertisements on the recruiter side"),
     ("Actors", "Recruiter, Administrator"),
     ("Entry criteria", "Recruiter validated by the administration"),
     ("Exit criteria", "Advertisement CRUD demonstrated")],
    [("US3.1", "As a recruiter, I log in to my space", "High", "2", "Done"),
     ("US3.2", "As a recruiter, I view my dashboard", "High", "2", "Done"),
     ("US3.3", "As a recruiter, I publish an advertisement", "High", "3", "Done"),
     ("US3.4", "As a recruiter, I edit an advertisement", "High", "2", "Done"),
     ("US3.5", "As a recruiter, I delete an advertisement", "Medium", "2", "Done"),
     ("US3.6", "As a recruiter, I view the applications", "High", "3", "Done"),
     ("US3.7", "As a recruiter, I view the advertisement history", "Medium", "2", "Done"),
     ("US3.8", "As a recruiter, I view a candidate's CV", "Medium", "2", "Done")],
    ["An unvalidated recruiter cannot publish an offer.",
     "Deletion of an offer is possible only with a confirmation step.",
     "A newly published offer has the 'pending' status until moderation.",
     "The recruiter sees only their own advertisements.",
     "Each action (create/edit/delete) is logged in the history.",
     "The dashboard summarizes the number of offers and applications."],
    [["Recruiter", "Authenticated, validated", "Publishes and manages advertisements"],
     ["Candidate", "Indirect actor", "Applies to the published offers"],
     ["Administrator", "Moderator", "Validates the recruiter and the offers"]],
    [["Use case name", "Publish an advertisement"],
     ["Actor", "Recruiter"],
     ["Precondition", "The recruiter is authenticated and validated"],
     ["Main flow",
      "1. The recruiter opens the publication form.\n2. The recruiter fills in the fields (title, description, sector, city, country).\n3. The controller validates the inputs.\n4. The DAL inserts the offer with status 'pending'.\n5. The view redisplays the dashboard with the new offer."],
     ["Alternative flows", "A1: a mandatory field is missing -> the form is redisplayed with errors.\nA2: the recruiter is not validated -> a message explains the situation."],
     ["Postcondition", "The offer exists with status 'pending' and appears in the dashboard"]],
    "[ UML use case diagram: Publish, Edit, Delete, View applications, View history ]",
    "[ UML class diagram: classes Recruiter, Advertisement, Application, History ]",
    "'Publish an advertisement' sequence diagram",
    "The authenticated recruiter fills in the advertisement form, the controller validates the fields, the DAL inserts "
    "the advertisement with status 'pending', the view redisplays the dashboard with the new advertisement.",
    "'View received applications' sequence diagram",
    "The recruiter selects an advertisement, the controller retrieves the related applications via the DAL, the view "
    "displays the list of candidates with the application date and a link to open the uploaded CV.",
    [("RecruiterId", "int", "Recruiter identifier (PK)"),
     ("CompanyName", "varchar(150)", "Company name"),
     ("Email", "varchar(100)", "Recruiter email"),
     ("IsVerified", "bit", "Validation state by the administration"),
     ("AdvertisementId", "int", "Advertisement identifier (PK)"),
     ("Title", "varchar(150)", "Advertisement title"),
     ("Description", "text", "Advertisement description"),
     ("Status", "varchar(20)", "pending / validated / rejected / deleted"),
     ("ApplicationId", "int", "Application identifier (PK)"),
     ("ApplicationDate", "datetime", "Application date"),
     ("ActionDate", "datetime", "History action date"),
     ("ActionType", "varchar(30)", "create / edit / delete")],
    ["The implemented interfaces include the recruiter dashboard, the advertisement creation/editing forms, the "
     "advertisement list with actions, the detail page and the page of candidates who applied.",
     "The dashboard presents a summary of the recruiter's activity (number of offers, number of applications) and quick "
     "links to the main actions. The list of advertisements uses a confirm dialog before deletion."],
    "[ Screenshot: recruiter dashboard and advertisement list ]",
    "[ Screenshot: advertisement creation form and applications page ]",
    "This sprint equipped recruiters with a complete space for managing advertisements and tracking applications. The "
    "CRUD operations on advertisements and the browsing of applications are operational.")

# ---- Sprint 4
sprint(doc,
    "Sprint 4: Administration space",
    [
        "This sprint covers the administration back-office: management of recruiters, moderation of offers, deletion of "
        "content and overall supervision of the platform. It is the control tower of JobzFactory and guarantees the "
        "quality and the legality of the published content.",
        "The sprint introduces the Administrator and OfferAction entities and the moderation methods of the JF.DAL layer. "
        "Each moderation action is logged with its type, its date and its author, which provides a full traceability of "
        "the content steering.",
        "The validation of the sprint is based on the demonstration of the validation of a recruiter, the moderation of "
        "a pending offer and the deletion of non-compliant content, with the corresponding log entries.",
    ],
    [("Sprint", "Sprint 4"), ("Release", "Release 1"), ("Duration", "1 week"),
     ("Objective", "Administer and moderate the platform"),
     ("Actors", "Administrator"),
     ("Entry criteria", "Recruiters and pending offers available"),
     ("Exit criteria", "Moderation scenarios demonstrated")],
    [("US4.1", "As an admin, I log in to the back-office", "High", "2", "Done"),
     ("US4.2", "As an admin, I view the admin dashboard", "High", "2", "Done"),
     ("US4.3", "As an admin, I validate a recruiter", "High", "2", "Done"),
     ("US4.4", "As an admin, I reject a recruiter", "Medium", "2", "Done"),
     ("US4.5", "As an admin, I moderate an offer", "High", "2", "Done"),
     ("US4.6", "As an admin, I delete an offer", "Medium", "2", "Done"),
     ("US4.7", "As an admin, I delete a recruiter", "Medium", "2", "Done"),
     ("US4.8", "As an admin, I view the list of offers", "Medium", "2", "Done"),
     ("US4.9", "As an admin, I view the moderation log", "Low", "2", "Done")],
    ["The administrator can activate or deactivate a recruiter.",
     "Any offer can be rejected or deleted by the administrator.",
     "Each moderation action is logged (type, date, author).",
     "A confirmation is required before any destructive action.",
     "The back-office is accessible only to the administrator role.",
     "The moderation log is immutable."],
    [["Administrator", "Super-user", "Validates and moderates the content"],
     ["Recruiter", "Indirect actor", "Is validated or rejected by the admin"],
     ["Offer", "Indirect entity", "Is moderated or deleted"]],
    [["Use case name", "Moderate an offer"],
     ["Actor", "Administrator"],
     ["Precondition", "The administrator is authenticated; a pending offer exists"],
     ["Main flow",
      "1. The admin opens the offer moderation list.\n2. The admin selects a pending offer.\n3. The admin chooses 'Validate' or 'Reject'.\n4. The controller updates the status via the DAL.\n5. The DAL logs the OfferAction.\n6. The view refreshes the list with the new status."],
     ["Alternative flows", "A1: the admin chooses 'Delete' -> the offer is marked deleted after confirmation."],
     ["Postcondition", "The offer has a new status and an OfferAction entry is logged"]],
    "[ UML use case diagram: Validate recruiter, Moderate offer, Delete content, View log ]",
    "[ UML class diagram: classes Administrator, Recruiter, Offer, OfferAction ]",
    "'Moderate an offer' sequence diagram",
    "The administrator selects a pending offer, the controller calls the DAL to update the status, the DAL logs the "
    "OfferAction, the view refreshes the list of offers with the new status.",
    "'Validate a recruiter' sequence diagram",
    "The administrator opens the recruiter list, selects a pending recruiter, the controller updates the IsVerified flag "
    "via the DAL, the recruiter is now able to publish offers, the view refreshes the list.",
    [("AdminId", "int", "Administrator identifier (PK)"),
     ("Email", "varchar(100)", "Administrator email"),
     ("OfferActionId", "int", "Moderation action identifier (PK)"),
     ("OfferId", "int", "Targeted offer (FK)"),
     ("AdminId", "int", "Author of the action (FK)"),
     ("ActionType", "varchar(30)", "Validation / Rejection / Deletion"),
     ("ActionDate", "datetime", "Action date"),
     ("Comment", "varchar(255)", "Optional moderation comment"),
     ("RecruiterId", "int", "Recruiter concerned (FK, nullable)"),
     ("IsVerified", "bit", "Recruiter validation state"),
     ("Status", "varchar(20)", "Offer status after the action"),
     ("LogImmutable", "bit", "Indicates that the log entry cannot be modified")],
    ["The implemented interfaces include the administrator dashboard, the recruiter list with validation actions, the "
     "offer list with moderation, and the detail and deletion pages with confirmation.",
     "The dashboard presents global counters (number of recruiters, number of pending offers, number of applications) "
     "and shortcuts to the moderation lists. The moderation log is accessible from a dedicated page."],
    "[ Screenshot: administrator dashboard and recruiter validation list ]",
    "[ Screenshot: offer moderation list and moderation log ]",
    "This sprint set up the administration back-office. Recruiter management and offer moderation are operational and "
    "allow full steering of the platform. The moderation log ensures the traceability of the content steering.")

doc.add_page_break()

# ===================== CHAPTER 6: RELEASE 2 =====================
chapter_title(doc, "CHAPTER 6: « Release 2 » (Sprints 5 to 7)")
add_body(doc,
    "The second release groups the three sprints that enrich the platform with the core business features: online "
    "application with CV upload, advanced search with pagination, and dashboards and statistics. At the end of this "
    "release, the platform delivers the full expected user journey, from the offer search to the application and the "
    "activity steering.")

# ---- Sprint 5
sprint(doc,
    "Sprint 5: Applications and CV upload",
    [
        "This sprint covers the application of candidates to offers, with the upload of a PDF CV, and the browsing of the "
        "application history. It is the heart of the business value of the platform: it transforms a visitor into an "
        "applicant.",
        "The sprint introduces the Application and CVFile entities and the file upload methods. The uploaded files are "
        "stored in a sandboxed folder with a per-offer and per-candidate organization, and only PDF files are accepted, "
        "with a maximum size enforced both client-side and server-side.",
        "The validation of the sprint is based on the demonstration of an application with CV upload, the uniqueness "
        "check (one application per offer per candidate) and the browsing of the application history.",
    ],
    [("Sprint", "Sprint 5"), ("Release", "Release 2"), ("Duration", "1 week"),
     ("Objective", "Enable online application"),
     ("Actors", "Candidate, Recruiter"),
     ("Entry criteria", "Candidate authenticated and offer validated"),
     ("Exit criteria", "Application + CV upload demonstrated")],
    [("US5.1", "As a candidate, I apply to an offer", "High", "3", "Done"),
     ("US5.2", "As a candidate, I upload my CV (PDF)", "High", "3", "Done"),
     ("US5.3", "As a candidate, I view my applications", "Medium", "2", "Done"),
     ("US5.4", "As a candidate, I cancel an application", "Low", "2", "Done"),
     ("US5.5", "As a recruiter, I view an uploaded CV", "Medium", "2", "Done"),
     ("US5.6", "As a system, I prevent double application", "High", "2", "Done"),
     ("US5.7", "As a system, I reject non-PDF uploads", "High", "2", "Done")],
    ["A candidate can apply only once to the same offer.",
     "Only PDF files are accepted for the CV, with a maximum size of 5 MB.",
     "The uploaded file is stored in a sandboxed, per-offer/per-candidate folder.",
     "The application date is set by the server, not by the client.",
     "The recruiter can open the CV only from the application list.",
     "The application history is sorted by date descending."],
    [["Candidate", "Authenticated user", "Applies and uploads the CV"],
     ["Recruiter", "Indirect actor", "Views the applications and the CVs"],
     ["File system", "Secondary actor", "Stores the uploaded PDF files"]],
    [["Use case name", "Apply to an offer"],
     ["Actor", "Candidate"],
     ["Precondition", "The candidate is authenticated; the offer is validated; the candidate has not yet applied"],
     ["Main flow",
      "1. The candidate opens the offer detail.\n2. The candidate clicks on 'Apply'.\n3. The candidate selects a PDF CV and submits.\n4. The controller checks the type and the size.\n5. The DAL records the application and stores the file.\n6. The view displays a confirmation."],
     ["Alternative flows",
      "A1: the file is not a PDF -> an error is displayed.\nA2: the candidate already applied -> a message informs them.\nA3: the file is too large -> an error is displayed."],
     ["Postcondition", "The application exists and the CV file is stored"]],
    "[ UML use case diagram: Apply, Upload CV, View applications, Cancel application ]",
    "[ UML class diagram: classes Application, Offer, Candidate, CVFile ]",
    "'Apply to an offer' sequence diagram",
    "The candidate opens the detail of an offer, clicks on 'Apply', uploads their PDF CV, the controller checks the "
    "type and size, the DAL records the application and stores the file, the view displays a confirmation.",
    "'View application history' sequence diagram",
    "The candidate opens the history page, the controller retrieves the applications of the candidate via the DAL, the "
    "view displays the list with the offer title, the application date and a link to the offer detail.",
    [("ApplicationId", "int", "Application identifier (PK)"),
     ("OfferId", "int", "Targeted offer (FK)"),
     ("CandidateId", "int", "Applying candidate (FK)"),
     ("CVPath", "varchar(255)", "Path of the uploaded CV file"),
     ("ApplicationDate", "datetime", "Application date"),
     ("Status", "varchar(20)", "Submitted / Cancelled"),
     ("CoverLetter", "text", "Optional cover letter"),
     ("FileSize", "int", "Size of the uploaded CV in bytes"),
     ("MimeType", "varchar(50)", "MIME type (application/pdf)"),
     ("CVFileId", "int", "CV file identifier (PK)"),
     ("OriginalName", "varchar(150)", "Original file name"),
     ("StoredName", "varchar(150)", "Server-side stored file name")],
    ["The implemented interfaces include the application form with a CV upload area, the confirmation page and the list "
     "of the candidate's application history. The upload area provides a drag-and-drop zone and displays the selected "
     "file name before submission.",
     "On the recruiter side, the application list shows the candidates who applied to an offer with a button to open the "
     "uploaded CV in a new tab. The history page is accessible from the candidate profile menu."],
    "[ Screenshot: application form with CV upload area ]",
    "[ Screenshot: application history and recruiter CV viewer ]",
    "This sprint established the business core of the platform: online application with CV upload. PDF file management, "
    "double-application prevention and application history are operational.")

# ---- Sprint 6
sprint(doc,
    "Sprint 6: Advanced search and pagination",
    [
        "This sprint covers the advanced search of offers by keyword, sector, city and country, with pagination of the "
        "results to guarantee acceptable performance even with a large volume of offers.",
        "The sprint introduces the SearchCriteria and ResultPage view-models and the paged query methods of the JF.DAL "
        "layer, based on the SQL Server OFFSET/FETCH clause. Indexed columns on the foreign keys and on the publication "
        "date ensure a stable response time.",
        "The validation of the sprint is based on the demonstration of a multi-criteria search with various filters and "
        "on the measurement of the response time on a sample of one thousand offers.",
    ],
    [("Sprint", "Sprint 6"), ("Release", "Release 2"), ("Duration", "1 week"),
     ("Objective", "Advanced search and pagination"),
     ("Actors", "Visitor, Candidate"),
     ("Entry criteria", "Sample data of ~1000 offers loaded"),
     ("Exit criteria", "Search + pagination demonstrated and timed")],
    [("US6.1", "As a visitor, I search by keyword", "High", "3", "Done"),
     ("US6.2", "As a visitor, I filter by sector", "High", "2", "Done"),
     ("US6.3", "As a visitor, I filter by city", "High", "2", "Done"),
     ("US6.4", "As a visitor, I filter by country", "Medium", "2", "Done"),
     ("US6.5", "As a visitor, I paginate the results", "Medium", "3", "Done"),
     ("US6.6", "As a visitor, I clear the filters", "Low", "1", "Done"),
     ("US6.7", "As a system, I index the search columns", "Medium", "2", "Done")],
    ["The search is performed on the title and the description of the offers.",
     "The number of offers per page is configurable (default 10).",
     "The search returns only validated offers.",
     "Filters are combined with a logical AND.",
     "The pagination uses OFFSET/FETCH for stable performance.",
     "An empty result set displays a clear 'no offer' message."],
    [["Visitor", "Unauthenticated user", "Enters criteria and browses the results"],
     ["Candidate", "Authenticated user", "Same as visitor, plus can apply"],
     ["Database", "Persistence service", "Executes the paged filtered query"]],
    [["Use case name", "Search for an offer"],
     ["Actor", "Visitor / Candidate"],
     ["Precondition", "Validated offers exist in the database"],
     ["Main flow",
      "1. The visitor opens the search page.\n2. The visitor enters a keyword and/or selects filters.\n3. The visitor submits the search.\n4. The controller forwards the criteria to the DAL.\n5. The DAL builds a filtered, paged query.\n6. The view displays the results and the pagination controls."],
     ["Alternative flows", "A1: no result -> a 'no offer' message is displayed.\nA2: invalid page index -> the first page is displayed."],
     ["Postcondition", "A page of results matching the criteria is displayed"]],
    "[ UML use case diagram: Search, Filter, Paginate, Clear filters ]",
    "[ UML class diagram: classes Offer, SearchCriteria, ResultPage, ActivitySector, City ]",
    "'Search for an offer' sequence diagram",
    "The visitor enters criteria, the controller forwards to the DAL which builds a filtered query with pagination "
    "(OFFSET/FETCH), the view displays the results and the pagination controls.",
    "'Paginate the results' sequence diagram",
    "The visitor clicks on the next page, the controller increments the page index, the DAL executes the paged query "
    "with the new offset, the view redisplays the results and updates the active page in the pagination controls.",
    [("Keyword", "varchar(100)", "Search term"),
     ("SectorId", "int", "Sector filter"),
     ("CityId", "int", "City filter"),
     ("CountryId", "int", "Country filter"),
     ("Page", "int", "Current page number"),
     ("PageSize", "int", "Number of offers per page"),
     ("TotalCount", "int", "Total number of matching offers"),
     ("TotalPages", "int", "Total number of pages"),
     ("Sort", "varchar(20)", "Sort criterion (date / title)"),
     ("OfferId", "int", "Offer identifier (indexed)"),
     ("PublicationDate", "datetime", "Indexed for sorting"),
     ("Title", "varchar(150)", "Indexed for keyword search")],
    ["The implemented interfaces include an advanced search panel with filters and a paginated result list with "
     "previous/next controls and page numbers. The active filters are summarized above the results and can be cleared "
     "in a single click.",
     "The result cards display the title, the sector, the city and the publication date, with a link to the offer "
     "detail. The pagination control is disabled when there is only one page."],
    "[ Screenshot: advanced search panel with filters ]",
    "[ Screenshot: paginated results and pagination controls ]",
    "This sprint improved the users' search experience. The multi-criteria advanced search and the pagination of the "
    "results are operational and keep a stable response time on a thousand offers.")

# ---- Sprint 7
sprint(doc,
    "Sprint 7: Dashboards and statistics",
    [
        "This sprint covers the generation of dashboards and statistics for recruiters and administrators: number of "
        "offers, number of applications, active recruiters, offers by sector. It closes the platform with a steering "
        "capability.",
        "The sprint introduces the Statistic and Indicator view-models and the aggregation methods of the JF.DAL layer, "
        "based on GROUP BY queries. The indicators are computed on the fly from the operational data, which guarantees "
        "their freshness without requiring a batch process.",
        "The validation of the sprint is based on the demonstration of the administrator dashboard and the recruiter "
        "dashboard, with counters and ranked lists consistent with the operational data.",
    ],
    [("Sprint", "Sprint 7"), ("Release", "Release 2"), ("Duration", "1 week"),
     ("Objective", "Steering and statistics"),
     ("Actors", "Recruiter, Administrator"),
     ("Entry criteria", "Operational data available (offers, applications)"),
     ("Exit criteria", "Dashboards demonstrated and consistent")],
    [("US7.1", "As an admin, I view the global statistics", "High", "3", "Done"),
     ("US7.2", "As an admin, I see the number of offers by sector", "Medium", "2", "Done"),
     ("US7.3", "As an admin, I see the top recruiters", "Medium", "2", "Done"),
     ("US7.4", "As a recruiter, I see my applications count", "Medium", "2", "Done"),
     ("US7.5", "As a recruiter, I see my offers by status", "Low", "2", "Done"),
     ("US7.6", "As a system, I aggregate the data with GROUP BY", "Medium", "2", "Done")],
    ["Statistics are computed on the fly from operational data.",
     "Indicators are presented as counters and ranked lists.",
     "The dashboards are restricted to the administrator and the recruiter roles.",
     "The aggregation queries use GROUP BY and indexed columns.",
     "The dashboard refreshes on each visit (no caching stale data).",
     "The ranked lists are limited to the top 10 entries."],
    [["Administrator", "Super-user", "Views the global dashboard"],
     ["Recruiter", "Authenticated user", "Views their own dashboard"],
     ["Database", "Persistence service", "Executes the aggregation queries"]],
    [["Use case name", "Generate a dashboard"],
     ["Actor", "Administrator / Recruiter"],
     ["Precondition", "The actor is authenticated; operational data exists"],
     ["Main flow",
      "1. The actor opens the dashboard page.\n2. The controller calls the DAL aggregation methods.\n3. The DAL executes GROUP BY queries (by sector, by recruiter, by status).\n4. The view renders the counters and the ranked lists."],
     ["Alternative flows", "A1: no data -> the counters display zero and the lists display a 'no data' message."],
     ["Postcondition", "The dashboard is displayed with up-to-date indicators"]],
    "[ UML use case diagram: View global stats, View offers by sector, View top recruiters, View recruiter stats ]",
    "[ UML class diagram: classes Statistic, Indicator, Period, AggregationResult ]",
    "'Generate a dashboard' sequence diagram",
    "The administrator opens the dashboard, the controller calls the DAL which aggregates the data (counts by sector, "
    "by recruiter, by period), the view renders the indicators as counters and lists.",
    "'View offers by sector' sequence diagram",
    "The administrator selects the 'offers by sector' view, the controller calls the DAL which executes a GROUP BY "
    "SectorId query, the view renders a ranked list of sectors with the number of offers per sector.",
    [("TotalOffers", "int", "Total number of offers"),
     ("TotalApplications", "int", "Total number of applications"),
     ("TotalRecruiters", "int", "Number of active recruiters"),
     ("TotalCandidates", "int", "Number of candidates"),
     ("OffersBySector", "table", "Distribution of offers by sector"),
     ("OffersByStatus", "table", "Distribution of offers by status"),
     ("TopRecruiters", "table", "Top recruiters by number of offers"),
     ("ApplicationsByPeriod", "table", "Applications grouped by period"),
     ("IndicatorLabel", "varchar(80)", "Label of the indicator"),
     ("IndicatorValue", "int", "Value of the indicator"),
     ("Rank", "int", "Rank in the ranked list"),
     ("PeriodStart", "datetime", "Start of the aggregation period")],
    ["The implemented interfaces include an administrator dashboard with counters (offers, applications, recruiters) "
     "and ranked lists (offers by sector, top recruiters), as well as a recruiter dashboard with the number of "
     "applications and the distribution of the recruiter's offers by status.",
     "The counters are displayed as highlighted cards at the top of the page, and the ranked lists are displayed as "
     "tables below. The layout is responsive and printable."],
    "[ Screenshot: administrator dashboard with counters and ranked lists ]",
    "[ Screenshot: recruiter dashboard and offers-by-status view ]",
    "This sprint made steering indicators available. The administrator and recruiter dashboards are operational and "
    "provide a synthetic view of the activity, closing the platform with a complete steering capability.")

# ===================== CONCLUSIONS AND RECOMMENDATIONS =====================
chapter_title(doc, "Conclusions and Recommendations")
add_body(doc,
    "This final year project was an opportunity to design and develop JobzFactory, a web-based recruitment platform "
    "aimed at connecting job seekers and recruiters. The platform offers candidates a profile management space, offer "
    "browsing and search, as well as an online application module with CV upload. It provides recruiters with a "
    "dedicated space for managing advertisements and tracking applications, and includes an administration back-office "
    "for content moderation and steering.")
add_body(doc,
    "The project was carried out following an agile SCRUM methodology, broken down into two releases grouping seven "
    "sprints. Each sprint followed a complete cycle of analysis, design, implementation and validation, allowing a "
    "progressive and controllable delivery of features. The chosen technical architecture, based on ASP.NET MVC, C# "
    "and SQL Server, proved robust and adaptable to the evolving needs of the project.")
add_body(doc,
    "Among the perspectives for evolution, we can mention: enriching the candidate profile with structured skills and "
    "experiences, setting up a matching engine between profiles and offers, sending email notifications when a new "
    "offer matches a profile, developing a mobile version of the application, as well as integrating automated testing "
    "and continuous integration tools to make future evolutions more reliable.")
add_body(doc,
    "This project was an opportunity to put into practice the knowledge acquired during the Professional Master's degree "
    "in Software Engineering, both in terms of software design and agile project management, and to build a complete, "
    "coherent and deployable application.")
doc.add_page_break()

# ===================== GLOSSARY =====================
chapter_title(doc, "Glossary")
gloss = [
    ("Agile", "Family of project management methodologies advocating an iterative and adaptive approach."),
    ("ASP.NET MVC", "Microsoft web framework based on the Model-View-Controller separation."),
    ("Backlog", "Prioritized list of features to be developed in an agile project."),
    ("Controller", "Component in charge of handling requests and orchestrating model and view."),
    ("CLR", "Common Language Runtime: the .NET execution engine."),
    ("DAL", "Data Access Layer: the data access layer."),
    ("Deployment", "Action of putting the application into production on a server."),
    ("IIS", "Internet Information Services: the Microsoft web server."),
    ("Job board", "Online platform for publishing and browsing job offers."),
    ("MVC", "Architectural pattern separating model, view and controller."),
    ("OFFSET/FETCH", "SQL Server clause used to implement server-side pagination."),
    ("Razor", "View engine used by ASP.NET MVC (.cshtml files)."),
    ("Release", "Released version grouping several sprints."),
    ("SCRUM", "Agile framework structured into sprints with defined roles and ceremonies."),
    ("Sprint", "Short iteration (1 to 4 weeks) during which an increment is developed."),
    ("UML", "Unified Modeling Language: a graphical software modeling language."),
    ("User story", "Concise description of a requirement from the user's perspective."),
    ("View model", "Object shaped specifically for a view, distinct from the business model."),
]
make_table(doc, ["Term", "Definition"], gloss, widths=[Cm(4), Cm(11)])
doc.add_page_break()

# ===================== BIBLIOGRAPHY =====================
chapter_title(doc, "Bibliography")
add_para(doc, "Reference books and documents:", bold=True, color=NAVY)
for b in [
    "P. Roques, UML 2 in Practice, Editions Eyrolles.",
    "C. E. Leiserson, T. H. Cormen, Introduction to Algorithms, Dunod.",
    "Microsoft, Official ASP.NET MVC documentation, Microsoft Docs.",
    "R. Pressman, Software Engineering: A Practitioner's Approach, McGraw-Hill.",
    "K. Schwaber, J. Sutherland, The Scrum Guide, Scrum.org.",
    "J. Richter, CLR via C#, Microsoft Press.",
    "M. Fowler, Patterns of Enterprise Application Architecture, Addison-Wesley.",
]:
    bullet(doc, b)
doc.add_page_break()

# ===================== WEBography =====================
chapter_title(doc, "Webography")
for w in [
    "https://learn.microsoft.com/aspnet/mvc",
    "https://learn.microsoft.com/dotnet/csharp",
    "https://learn.microsoft.com/sql/ssms",
    "https://learn.microsoft.com/iis",
    "https://www.scrum.org/resources/scrum-guide",
    "https://getbootstrap.com/",
    "https://developer.mozilla.org/",
    "https://www.uml.org/",
    "https://learn.microsoft.com/dotnet/framework/data/adonet/",
]:
    bullet(doc, w)
doc.add_page_break()

# ===================== APPENDIX A =====================
chapter_title(doc, "Appendix A: Background on recruitment platforms")
add_body(doc,
    "This appendix presents a theoretical reminder on online recruitment platforms (job boards) and on the ASP.NET MVC "
    "framework, in order to provide the reader with the conceptual context necessary to understand the project.")
h2(doc, "1. Online recruitment platforms")
add_body(doc,
    "An online recruitment platform, or job board, is a website that allows companies to publish their job offers and "
    "job seekers to browse these offers and apply. These platforms have become a major channel of modern recruitment, "
    "because of their reach, their immediacy and their ability to centralize applications.")
add_body(doc,
    "The features expected of a modern job board include: the management of differentiated user accounts (candidates, "
    "recruiters, administrators), the publication and moderation of offers, multi-criteria search, CV upload, application "
    "tracking and the production of statistical indicators.")
h2(doc, "2. The ASP.NET MVC framework")
add_body(doc,
    "ASP.NET MVC is a Microsoft web framework based on the Model-View-Controller architectural pattern. It separates "
    "concerns: the model carries the data and the business logic, the view generates the user interface, and the "
    "controller orchestrates the interactions. This separation improves maintainability, testability and collaboration.")
add_body(doc,
    "ASP.NET MVC relies on the Razor view engine (.cshtml), a declarative URL routing system and a tight integration "
    "with the .NET platform and the C# language. It is particularly well suited to the development of professional web "
    "applications structured in layers.")
h2(doc, "3. The SCRUM methodology")
add_body(doc,
    "SCRUM is an agile framework that structures the project into short sprints, framed by ceremonies (sprint planning, "
    "daily scrum, sprint review, retrospective) and roles (Product Owner, Scrum Master, development team). The product "
    "is described by a prioritized backlog, and each sprint delivers a potentially usable increment.")
add_body(doc,
    "SCRUM favors adaptation to change, transparency and continuous improvement. It was chosen for this project because "
    "of the evolving nature of the requirements and the need to deliver regular functional increments, evaluable at the "
    "end of each sprint.")
doc.add_page_break()

# ===================== APPENDIX B: CODE EXTRACT =====================
chapter_title(doc, "Appendix B: Code extract (representative controller method)")
add_body(doc,
    "This appendix presents an excerpt of the OffreController of the public front-office. The Index action retrieves the "
    "paginated list of validated offers by delegating to the JF.DAL layer, then returns a Razor view with the result page. "
    "The code is shown in a monospaced block for readability.")
add_code(doc,
'''public class OffreController : Controller
{
    private readonly OffreRepository _repo = new OffreRepository();

    // GET: /Offre/Index?page=1&keyword=...
    public ActionResult Index(int page = 1, string keyword = "",
        int? sectorId = null, int? cityId = null)
    {
        const int pageSize = 10;
        if (page < 1) page = 1;

        SearchCriteria criteria = new SearchCriteria
        {
            Keyword = keyword,
            SectorId = sectorId,
            CityId = cityId,
            Page = page,
            PageSize = pageSize,
            // Only validated offers are publicly visible (business rule BR-07)
            Status = "validated"
        };

        ResultPage<Offre> result = _repo.SearchPaged(criteria);

        ViewBag.Keyword = keyword;
        ViewBag.SectorId = sectorId;
        ViewBag.CityId = cityId;
        ViewBag.TotalPages = result.TotalPages;
        ViewBag.CurrentPage = page;

        return View(result.Items);
    }

    // GET: /Offre/Details/5
    public ActionResult Details(int id)
    {
        Offre offre = _repo.GetById(id);
        if (offre == null || offre.Status != "validated")
            return HttpNotFound("Offer not found.");

        return View(offre);
    }
}''', caption=None)
add_body(doc,
    "The two actions illustrate the separation of responsibilities: the controller validates the inputs, builds the "
    "criteria, delegates the data access to the repository and selects the view. The business rule that restricts the "
    "public display to validated offers is enforced at the data access level through the criteria's Status field.")
doc.add_page_break()

# ===================== APPENDIX C: DATABASE SCHEMA =====================
chapter_title(doc, "Appendix C: Database schema script excerpt")
add_body(doc,
    "This appendix presents an excerpt of the database creation script. It defines the main tables (Offre, Recruteur, "
    "Profil, Offre_Postuler) and the foreign keys ensuring the referential integrity between the offers, the recruiters, "
    "the candidates and the references (sectors, cities, countries). The script is shown in a monospaced block.")
add_code(doc,
'''CREATE TABLE dbo.Recruteur (
    RecruteurId     INT IDENTITY(1,1) NOT NULL,
    CompanyName     VARCHAR(150) NOT NULL,
    Email           VARCHAR(100) NOT NULL,
    PasswordHash    VARCHAR(255) NOT NULL,
    IsVerified      BIT NOT NULL DEFAULT 0,
    CreatedAt       DATETIME NOT NULL DEFAULT GETDATE(),
    CONSTRAINT PK_Recruteur PRIMARY KEY (RecruteurId),
    CONSTRAINT UQ_Recruteur_Email UNIQUE (Email)
);

CREATE TABLE dbo.Offre (
    OffreId         INT IDENTITY(1,1) NOT NULL,
    Title           VARCHAR(150) NOT NULL,
    Description     TEXT NOT NULL,
    SectorId        INT NOT NULL,
    CityId          INT NOT NULL,
    CountryId       INT NOT NULL,
    RecruiterId     INT NOT NULL,
    PublicationDate DATETIME NOT NULL DEFAULT GETDATE(),
    Status          VARCHAR(20) NOT NULL DEFAULT 'pending',
    CONSTRAINT PK_Offre PRIMARY KEY (OffreId),
    CONSTRAINT FK_Offre_Secteur  FOREIGN KEY (SectorId)
        REFERENCES dbo.SecteurActivite(SecteurId),
    CONSTRAINT FK_Offre_Ville    FOREIGN KEY (CityId)
        REFERENCES dbo.Ville(VilleId),
    CONSTRAINT FK_Offre_Pays     FOREIGN KEY (CountryId)
        REFERENCES dbo.Pays(PaysId),
    CONSTRAINT FK_Offre_Recruteur FOREIGN KEY (RecruiterId)
        REFERENCES dbo.Recruteur(RecruteurId),
    CONSTRAINT CK_Offre_Status CHECK
        (Status IN ('pending','validated','rejected','deleted'))
);

CREATE TABLE dbo.Profil (
    ProfilId        INT IDENTITY(1,1) NOT NULL,
    Email           VARCHAR(100) NOT NULL,
    PasswordHash    VARCHAR(255) NOT NULL,
    LastName        VARCHAR(50) NOT NULL,
    FirstName       VARCHAR(50) NOT NULL,
    Phone           VARCHAR(20) NULL,
    RegistrationDate DATETIME NOT NULL DEFAULT GETDATE(),
    AccountState    VARCHAR(20) NOT NULL DEFAULT 'pending',
    CONSTRAINT PK_Profil PRIMARY KEY (ProfilId),
    CONSTRAINT UQ_Profil_Email UNIQUE (Email)
);

CREATE TABLE dbo.Offre_Postuler (
    ApplicationId   INT IDENTITY(1,1) NOT NULL,
    OffreId         INT NOT NULL,
    ProfilId        INT NOT NULL,
    CVPath          VARCHAR(255) NOT NULL,
    ApplicationDate DATETIME NOT NULL DEFAULT GETDATE(),
    Status          VARCHAR(20) NOT NULL DEFAULT 'submitted',
    CONSTRAINT PK_OffrePostuler PRIMARY KEY (ApplicationId),
    CONSTRAINT FK_App_Offre  FOREIGN KEY (OffreId)
        REFERENCES dbo.Offre(OffreId),
    CONSTRAINT FK_App_Profil FOREIGN KEY (ProfilId)
        REFERENCES dbo.Profil(ProfilId),
    CONSTRAINT UQ_App_Unique UNIQUE (OffreId, ProfilId) -- one application per offer
);''', caption=None)
add_body(doc,
    "The UNIQUE constraint on (OffreId, ProfilId) implements the business rule BR-03 at the database level: a candidate "
    "can apply only once to a given offer. The CHECK constraint on the Offre status implements the status enumeration "
    "and prevents any inconsistent value.")
doc.add_page_break()

# ===================== APPENDIX D: EXTRA SCREENSHOTS =====================
chapter_title(doc, "Appendix D: Additional screenshots")
add_body(doc,
    "This appendix gathers additional screenshots of the implemented platform, complementing the figures presented in "
    "the sprint chapters. They illustrate the rendering of the main interfaces on desktop and tablet screens.")
add_figure(doc, "Additional screenshot - public home page (desktop)",
           "[ Screenshot: JobzFactory public home page on a desktop screen ]", box_h=Cm(8))
add_figure(doc, "Additional screenshot - offer detail (tablet)",
           "[ Screenshot: offer detail page on a tablet screen ]", box_h=Cm(8))
add_figure(doc, "Additional screenshot - recruiter dashboard",
           "[ Screenshot: recruiter dashboard with counters and lists ]", box_h=Cm(8))
add_figure(doc, "Additional screenshot - administration back-office",
           "[ Screenshot: administration back-office moderation list ]", box_h=Cm(8))
add_figure(doc, "Additional screenshot - application form with CV upload",
           "[ Screenshot: application form with drag-and-drop CV upload ]", box_h=Cm(8))
add_figure(doc, "Additional screenshot - advanced search results",
           "[ Screenshot: advanced search results with active filters and pagination ]", box_h=Cm(8))

# ----------------------------- populate front matter lists -----------------------------
fill_list_before(lof_marker, FIG_LIST)
fill_list_before(lot_marker, TAB_LIST)

# ----------------------------- save -----------------------------
os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print("Report generated:", OUT)
print("Figures total:", FIG)
print("Tables total:", TAB)
